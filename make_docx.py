from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Style helpers ─────────────────────────────────────────────────────────
BLUE  = RGBColor(0x25, 0x63, 0xEB)
PURP  = RGBColor(0x7C, 0x3A, 0xED)
DARK  = RGBColor(0x1E, 0x29, 0x3B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGREY = RGBColor(0xF1, 0xF5, 0xF9)
DKBLU = RGBColor(0x1E, 0x3A, 0x5F)

def set_cell_bg(cell, rgb_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)

def cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DSAI4205 — Big Data Analytics\nExam Review Notes')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DKBLU
    doc.add_paragraph('PolyU · Dr. Ken Fong · Lectures 1–10 · Exam = 40% of grade\n'
                       'Format: Multiple Choice + Long Questions').alignment \
        == WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f'  {text}  ')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = WHITE
    # shade paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2563EB')
    pPr.append(shd)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = PURP

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.add_run(text).font.size = Pt(9.5)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.add_run(text).font.size = Pt(9.5)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1E293B')
    pPr.append(shd)

def callout(label, text, colour='DBEAFE', text_colour=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f'{label}  {text}')
    run.font.size = Pt(9)
    if text_colour:
        run.font.color.rgb = text_colour
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), colour)
    pPr.append(shd)

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '2563EB')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(8.5)
        run.font.color.rgb = WHITE
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = t.rows[r+1].cells[c]
            if (r % 2) == 1:
                set_cell_bg(cell, 'F8FAFC')
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def pg():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
cover()

# ─────────────────────────────────────────────────────────────────────────
h1('L1 · Introduction to Big Data Analytics')

h2('The 6 V\'s of Big Data')
table(
    ['V', 'Definition', 'Example'],
    [
        ['Volume',      'Data amount — growing exponentially',                           '6.4B smartphone subscriptions (2022)'],
        ['Velocity',    'Speed of generation AND processing',                            '16.4B Google searches/day; real-time stock ticks'],
        ['Variety',     'Diversity of types (structured, semi-structured, unstructured)','Text, video, sensor logs, JSON, CSV simultaneously'],
        ['Veracity',    'Quality, accuracy, trustworthiness of data',                   'Measurement errors, human errors, inconsistent sources'],
        ['Value',       'Business usefulness; longer processing = less value',           'Not all collected data has business worth'],
        ['Variability', 'How fast/extent data changes; same word, multiple meanings',   'Netflix prefs shift when new releases drop'],
    ],
    [2.2, 6.5, 6.5]
)

h2('Types of Analytics')
table(
    ['Type', 'Question', 'Example'],
    [
        ['Descriptive',  'What happened?',       'Sales report last quarter'],
        ['Diagnostic',   'Why did it happen?',   'Root-cause analysis of churn'],
        ['Predictive',   'What will happen?',    'Demand forecasting'],
        ['Prescriptive', 'What should we do?',   'Route optimisation recommendation'],
    ],
    [2.5, 4, 8.7]
)

callout('★ Netflix case study:', 'Map 4 Vs to Netflix challenges (Volume→Spark storage, Velocity→Spark Streaming, Variety→SparkSQL, Veracity→validation+ML)', 'DBEAFE')

# ─────────────────────────────────────────────────────────────────────────
h1('L2 · The Path to Parallelism')

h2('MapReduce Steps')
for i, s in enumerate([
    'Input: bag of (key, value) pairs',
    'Map: each worker processes chunk, emits intermediate (key, value) pairs',
    'Shuffle: group all intermediate pairs by key → same reducer',
    'Reduce: aggregate values per key into final result',
    'Output: write final (key, value) pairs to storage',
], 1):
    bullet(f'Step {i}: {s}')

h2('Word Count (classic MapReduce)')
code('Map:    ("doc", "hello world hello") → [("hello",1),("world",1),("hello",1)]\nShuffle: ("hello",[1,1]), ("world",[1])\nReduce: ("hello",2), ("world",1)')

h2('HDFS')
bullet('Fixed-size block partitioning (default 128 MB)')
bullet('3-way replication across 3 separate nodes (fault tolerance)')
bullet('Write-once, read-many — files are immutable once written')
bullet('NameNode: master (stores metadata); DataNode: worker (stores blocks)')

h2('MapReduce vs Spark')
table(
    ['MapReduce', 'Spark'],
    [
        ['Writes intermediate to disk (slow)',      'Keeps data in memory (10–100× faster for iterative)'],
        ['Only Map + Reduce primitives',            'Rich 100+ operation library'],
        ['No iterative algorithm support',          'Efficient for ML, graph, streaming'],
    ],
    [8, 7.2]
)

h2('MapReduce — Detailed 3-Mapper Walkthrough')
code('3 documents across 3 mappers:\n  Mapper 1: "I often repeat repeat"\n  Mapper 2: "I do not not"\n  Mapper 3: "repeat after me"\n\nMap phase (each emits (word,1) pairs):\n  Mapper 1: (I,1)(often,1)(repeat,1)(repeat,1)\n  Mapper 2: (I,1)(do,1)(not,1)(not,1)\n  Mapper 3: (repeat,1)(after,1)(me,1)\n\nShuffle (group by key, sent across network):\n  (I,[1,1])  (repeat,[1,1,1])  (not,[1,1])  (often,[1])\n\nReduce (sum list):\n  (I,2)  (repeat,3)  (not,2)  (often,1)\n\nPerformance cost: ALL intermediate data written to HDFS disk between phases\n                  (needed for fault tolerance — but makes iterative ML slow)\nSpark fix: keep intermediate data in MEMORY → 10-100x faster')

h2('Dask — Block Algorithms & Task Graph')
code('Block algorithm example (4x4 array split into 2x2 chunks):\n  Step 1: each worker computes sum of its 2x2 chunk independently\n  Step 2: combine chunk results → total\n  Worker = distinct server / local machine / CPU core\n\nTask Graph:\n  - Symbolic representation of computations\n  - Nodes = tasks; Edges = dependencies\n  - Built lazily; executed only on .compute()\n  - Visualise: result.visualize()  (uses Graphviz)')

h2('Dask Chunk Specification (4 ways)')
code('# 1. Same size for all dimensions\narr = da.ones((10000, 12000), chunks=1000)\n\n# 2. Different size per dimension\narr = da.ones((10000, 12000), chunks=(5000, 4000))\n\n# 3. Dict per dimension\narr = da.ones((10000, 12000), chunks={0: 1000, 1: 2000})\n\n# 4. Fully explicit block sizes\narr = da.ones((10000, 12000), chunks=((5000,5000),(4000,4000,4000,4000)))\n\n# Rechunk (EXPENSIVE!)\narr.rechunk({0: -1, 1: 100})  # -1 = keep entire dim as one chunk\n# WARNING: bad chunk choice → 86ms vs 7min 13s!')
callout('Chunk size guideline:', '100 MB – 1 GB per chunk. Too small = task graph overhead. Too large = out-of-memory.', 'FFFBEB')

h2('Dask Distributed Client & persist()')
code('from dask.distributed import Client\nclient = Client(n_workers=4)   # local cluster\n# Dashboard: http://127.0.0.1:8787/status\n\n# persist() — cache intermediate result in distributed memory\nddf_jfk = ddf[ddf[\'Origin\'] == \'JFK\']\nddf_jfk = ddf_jfk.persist()   # stays in RAM for multiple downstream ops')

h2('Zarr — Chunked Array Storage')
code('arr.to_zarr(\'my_array.zarr\')      # save to chunked binary format\narr2 = da.from_zarr(\'my_array.zarr\')  # load back')

# ─────────────────────────────────────────────────────────────────────────
pg()
h1('L3 · Apache Spark — RDDs & DataFrames')

h2('RDD — Data Lineage & Fault Tolerance')
code('Problem: a node fails mid-computation — how to recover?\n\nRDD solution — Data Lineage:\n  Record OPERATIONS (not data) applied to each partition\n  Each RDD knows: "I was derived from parent X by operation F"\n  Recovery: re-run F on surviving partitions of X → rebuild lost partition\n  Cost: just replay operations; no expensive data replication needed\n\nImmutability: cannot modify an RDD → always create a new one\n  Guarantees deterministic lineage (no side effects)')

h2('RDD Granularity of Data Flow')
table(
    ['Mode', 'Description', 'When to use'],
    [
        ['Whole Dataset', 'Entire dataset through one transformation at once', 'Small data that fits in memory'],
        ['Row',           'Each row processed independently',                   'Flexible but slow (high overhead per task)'],
        ['Partition ✓',   'User-defined partition count; each worker handles one', 'Recommended — balances parallelism and overhead'],
    ],
    [3, 6.5, 5.7]
)
code('rdd.getNumPartitions()    # check current partition count\nrdd.repartition(10)       # increase partitions (triggers shuffle)\nrdd.coalesce(2)           # reduce partitions (no shuffle)')

h2('Core Concepts')
table(
    ['Concept', 'Description'],
    [
        ['RDD',              'Resilient Distributed Dataset — immutable, distributed, fault-tolerant via lineage, lazy'],
        ['DataFrame',        'Named-column structure; optimised by Catalyst; lazy; supports StructType/ArrayType/MapType'],
        ['Lazy evaluation',  'Transformations build a DAG; nothing executes until an Action is called'],
        ['Data lineage',     'Record of all operations to rebuild an RDD — enables fault recovery without replication'],
        ['Immutability',     'Cannot modify an RDD; always produce a new one'],
    ],
    [3.5, 11.7]
)

h2('Transformations vs Actions')
table(
    ['Category', 'Description', 'Examples'],
    [
        ['Transformation (lazy)',  'Returns new RDD/DF; no compute yet', 'map, filter, flatMap, join, union, reduceByKey, groupByKey, distinct, sortBy, cartesian, subtract, intersection'],
        ['Action (eager)',         'Triggers execution; returns value or writes', 'collect, count, show, save, first, take(n), takeOrdered, min, max, sum, mean, stdev'],
    ],
    [3.5, 5, 6.7]
)

h2('Narrow vs Wide Transformations')
table(
    ['', 'Narrow', 'Wide (Shuffle)'],
    [
        ['Mapping',    '1 input partition → 1 output partition', 'Multiple partitions → output partitions'],
        ['Network',    'No shuffle needed',                      'Network communication required'],
        ['Examples',   'filter, map, flatMap',                   'groupByKey, sortBy, join, reduceByKey'],
        ['Cost',       'Fast; pipelines stages',                 'Expensive; stage boundary in DAG'],
    ],
    [3, 6.5, 5.7]
)

h2('Key RDD Operations (quick reference)')
table(
    ['Op', 'What it does', 'Example result'],
    [
        ['map(f)',            'Apply f to every element',                       '[1,2,3] → [2,3,4]'],
        ['flatMap(f)',        'Like map but flatten output (1 level)',           '["a b","c"] → ["a","b","c"]'],
        ['filter(f)',         'Keep elements where f is True',                  '[1,2,3] filter x>2 → [3]'],
        ['reduceByKey(f)',    'Aggregate values with same key',                 '[(a,1),(a,2)] → [(a,3)]'],
        ['mapValues(f)',      'Apply f to values only; key unchanged',          '[(a,2)] → [(a,4)]  (f=x²)'],
        ['flatMapValues(f)', 'Like mapValues but flattens list output into individual elements', '[(a,[1,2])] → [(a,1),(a,2)]'],
        ['groupByKey()',      'Group values by key into iterable',              '[(a,1),(a,2)] → [(a,[1,2])]'],
        ['join(rdd2)',        'Inner join on key',                              '[(3,a)] join [(3,b)] → [(3,(a,b))]'],
        ['leftOuterJoin()',   'All keys from left; None for missing right',     '[(1,a)] → [(1,(a,None))]'],
        ['subtractByKey()',   'Keys in left not in right',                      '[(1,a),(3,b)] − [(3,c)] → [(1,a)]'],
        ['cartesian()',       'All pairs (A × B)',                              '[1,2]×[3] → [(1,3),(2,3)]'],
        ['union()',           'Combine all elements (keeps dupes)',             '[1,2]∪[3] → [1,2,3]'],
        ['intersection()',   'Only shared elements',                           '[1,2,3]∩[2,3,4] → [2,3]'],
        ['distinct()',        'Remove duplicates',                              '[1,1,2] → [1,2]'],
        ['sortBy(f)',         'Sort by key function',                           'sortBy value descending'],
        ['randomSplit([w])',  'Split RDD by weights (e.g. train/test)',         '[0.8, 0.2]'],
        ['takeOrdered(n)',    'Smallest n elements (or custom key)',            '[10,1,2,9] → [1,2]'],
        ['collect()',         '★ ACTION: return all as Python list',           'triggers computation'],
        ['count()',           '★ ACTION: number of elements',                  '3'],
        ['first()',           '★ ACTION: first element',                       '21'],
        ['take(n)',           '★ ACTION: first n elements',                    '[21,1]'],
    ],
    [3.5, 6, 5.7]
)

h2('DataFrame Key Operations')
table(
    ['Op', 'SQL equivalent', 'Note'],
    [
        ['select(col)',              'SELECT col',      'Projection'],
        ['filter() / where()',       'WHERE',           'Row filtering'],
        ['groupBy().agg()',          'GROUP BY + agg',  'Aggregation'],
        ['join(df2, condition)',     'JOIN',            'Combine two DFs'],
        ['orderBy() / sort()',       'ORDER BY',        'Sort rows'],
        ['withColumn(name, expr)',   '(computed col)',  'Add or replace column'],
        ['drop(col)',                '—',               'Remove column'],
        ['printSchema()',            'DESCRIBE',        'Show column types'],
    ],
    [4, 3.5, 7.7]
)

h2('Catalyst Optimizer')
bullet('DataFrame operations build an Abstract Syntax Tree (AST) of the logical plan')
bullet('Catalyst rewrites plan: predicate pushdown, column pruning, join reordering')
bullet('Pushes .where() filters BEFORE joins → reads less data → faster')
callout('Why DataFrames > RDDs:', 'Catalyst optimises queries automatically. RDDs execute exactly as written.', 'ECFDF5')

h2('SparkSession Setup')
code('from pyspark.sql import SparkSession\nss = SparkSession.builder\\\n     .master("local[4]")   # local[x]: x = CPU threads\n     .appName("MyApp")\\\n     .getOrCreate()\nspark = ss.sparkContext  # for RDD operations')

h2('Nested Data Types in DataFrames')
table(
    ['Type', 'Description', 'Example'],
    [
        ['StructType', 'Named fields (sub-record)',  'address: {city:"HK", zip:"000"}'],
        ['ArrayType',  'List of values',             'tags: ["big","data"]'],
        ['MapType',    'Key-value map',              'attrs: {height:170, weight:65}'],
    ],
    [3, 6, 6.2]
)

h2('map() vs mapValues() vs flatMapValues()')
code('rdd = sc.parallelize([(\'a\',[1,2,3]),(\'b\',[4,5])])\n\n# map() — applied to ENTIRE element (key + value)\nrdd.map(lambda x: (x[0], len(x[1]))).collect()\n# [(\'a\',3),(\'b\',2)]\n\n# mapValues() — applied to VALUE only; key preserved\nrdd.mapValues(lambda v: sum(v)).collect()\n# [(\'a\',6),(\'b\',9)]\n\n# flatMapValues() — like mapValues but FLATTENS the list output\nrdd.flatMapValues(lambda v: v).collect()\n# [(\'a\',1),(\'a\',2),(\'a\',3),(\'b\',4),(\'b\',5)]  each item becomes its own row')

h2('More DataFrame Patterns')
code('# Create DataFrame from Python list\ndata = [(\'Alice\',30),(\'Bob\',25)]\ndf = spark.createDataFrame(data, [\'name\',\'age\'])\n\n# Rename columns with toDF()\nrdd.toDF([\'letter\',\'number\'])\n\n# Join two DataFrames\ndf1.join(df2, \'id\', \'inner\').show()\n\n# Sort descending\nfrom pyspark.sql.functions import col, desc\ndf.orderBy(col(\'price\').desc(), col(\'id\').desc()).show()\n\n# Unique values\ndf.select(\'category\').distinct().show()\n\n# DataFrame → RDD\nrdd = df.rdd\nrdd.map(lambda row: (row[\'name\'], row[\'age\'])).collect()\n\n# leftOuterJoin — handle None for missing values\njoined = rdd1.leftOuterJoin(rdd2)\nresult = joined.map(lambda x: (x[0], x[1][0], x[1][1] or \'MISSING\'))')

h2('Pandas — apply() & reset_index()')
code('# apply() row-wise function (axis=1)\ndf[\'full_name\'] = df.apply(lambda r: r[\'first\']+\' \'+r[\'last\'], axis=1)\n\n# reset_index() — convert grouped index back to column\ngrouped = df.groupby(\'dept\')[\'salary\'].mean().reset_index()\n# \'dept\' is now a regular column')

h2('Inverted Index (Take-Home Exercise)')
body('Classic IR structure: word → sorted list of document IDs containing it')
code('# Input: "D1 || word1 word2 word3"\nlines = spark.textFile("docs.txt")\ninv_index = lines\\\n  .flatMap(lambda line: [\n      (word, line.split("||")[0].strip())\n      for word in line.split("||")[1].strip().split()\n  ])\\\n  .distinct()\\\n  .groupByKey()\\\n  .mapValues(lambda docs: sorted(list(docs)))\\\n  .sortByKey()')

# ─────────────────────────────────────────────────────────────────────────
pg()
h1('L4 · Hive, Shark & SparkSQL')

h2('Data Formats Comparison')
table(
    ['Format', 'Type', 'Best for'],
    [
        ['CSV',     'Row, text',       'Human-readable; slow analytics'],
        ['JSON',    'Semi-structured', 'Flexible schema; verbose'],
        ['Parquet', 'Columnar, binary','Analytics — column pruning + compression (preferred)'],
        ['ORC',     'Columnar, binary','Hive; great compression'],
        ['Avro',    'Row, binary',     'Streaming + schema evolution'],
    ],
    [2.5, 3.5, 9.2]
)

h2('Hive vs SparkSQL')
table(
    ['', 'Hive', 'SparkSQL'],
    [
        ['Execution engine', 'MapReduce (slow)', 'Spark in-memory (fast)'],
        ['Latency',          'High (minutes)',   'Low (seconds)'],
        ['Iterative queries','Poor',             'Excellent'],
        ['Use case',         'Batch ETL',        'Interactive + batch analytics'],
    ],
    [3.5, 5.5, 6.2]
)

h2('HiveQL — Key Commands')
code('-- Load local file into Hive table\nLOAD DATA LOCAL INPATH \'/home/user/data.csv\' INTO TABLE sales;\n\n-- Load HDFS file into Hive table (moves the file)\nLOAD DATA INPATH \'/hdfs/data/sales.csv\' INTO TABLE sales;\n\n-- Load into a specific partition\nLOAD DATA INPATH \'/hdfs/data/jan.csv\'\nINTO TABLE sales PARTITION (year=2024, month=1);\n\n-- Create partitioned table\nCREATE TABLE logs (user_id STRING, action STRING)\nPARTITIONED BY (dt STRING) STORED AS PARQUET;\n\n-- Query single partition (avoids full scan)\nSELECT * FROM logs WHERE dt = \'2024-01-01\';')

h2('Hive Buckets (Bucketing)')
body('Partitioning divides data into directories by value. Bucketing further divides each partition into fixed files via hash on a column.')
code('Bucket assignment:  bucket_id = hash(column_value) % num_buckets\n\nBenefits:\n  - Efficient TABLESAMPLE queries\n  - Faster joins when both tables bucketed on join key (bucket map join)\n  - Avoids data skew vs pure range partitioning\n\nCREATE TABLE orders (order_id INT, customer STRING, amount DOUBLE)\nPARTITIONED BY (year INT)\nCLUSTERED BY (customer) INTO 32 BUCKETS\nSTORED AS ORC;')

h2('Data Lake vs Data Warehouse')
table(
    ['', 'Data Lake', 'Data Warehouse'],
    [
        ['Schema',      'Schema-on-read (flexible)',  'Schema-on-write (strict)'],
        ['Data',        'Raw, all formats',           'Cleaned, structured'],
        ['Ingestion',   'Fast, all data welcome',     'Slow (ETL required)'],
        ['Query speed', 'Slower (must infer schema)', 'Faster (pre-optimised)'],
        ['Risk',        '"Data swamp" without governance', 'Stale if ETL pipeline breaks'],
    ],
    [3, 5.5, 6.7]
)

# ─────────────────────────────────────────────────────────────────────────
h1('L5 · Natural Language Processing (NLP)')

h2('Linguistics Framework (NLP Subfields)')
table(
    ['Subfield', 'What it studies', 'Example'],
    [
        ['Phonetics',  'Physical sounds of speech',                   'IPA transcription, acoustic signals'],
        ['Phonology',  'Sound patterns and rules in a language',      '"cats" → /kæts/ not /kætz/'],
        ['Morphology', 'Word structure and formation',                '"un-help-ful" = prefix + root + suffix'],
        ['Syntax',     'Grammar rules for combining words',           'Subject-Verb-Object order in English'],
        ['Semantics',  'Meaning of words and sentences',              'Synonymy, antonymy, word sense disambiguation'],
        ['Pragmatics', 'Language use in context / speaker intent',   '"Can you pass the salt?" = request'],
    ],
    [2.8, 5, 7.4]
)

h3('Morphology: Inflectional vs Derivational')
table(
    ['Type', 'Description', 'Example'],
    [
        ['Inflectional', 'Changes grammatical form, same word class', 'dog→dogs (plural), walk→walked (past)'],
        ['Derivational', 'Creates new words, often changes class',    'happy→happiness (adj→noun), teach→teacher'],
    ],
    [3.5, 5.5, 6.2]
)

h2('NLP Applications')
for s in ['Sentiment Analysis: classify opinions as positive/negative/neutral',
          'Machine Translation: Google Translate, DeepL',
          'Chatbots & Virtual Assistants: Siri, Alexa, ChatGPT',
          'Speech Recognition: Whisper, Dragon NaturallySpeaking',
          'Named Entity Recognition (NER): extract persons, places, organisations',
          'Information Retrieval: search engines',
          'Text Summarisation: abstractive vs extractive']:
    bullet(s)

h2('NLTK Library Overview')
code('import nltk\nfrom nltk.tokenize import word_tokenize\nfrom nltk.corpus import stopwords\nfrom nltk.stem import PorterStemmer, WordNetLemmatizer\n\ntokens = word_tokenize("I am running fast.")\n# [\'I\',\'am\',\'running\',\'fast\',\'.\']\n\nsw = set(stopwords.words(\'english\'))\nfiltered = [w for w in tokens if w.lower() not in sw]\n\nps = PorterStemmer()\nps.stem(\'computational\')   # \'comput\'\n\nwnl = WordNetLemmatizer()\nwnl.lemmatize(\'better\', pos=\'a\')   # \'good\' (needs POS tag)')

h2('Text Normalisation — Regex Patterns')
code('import re\n\ndef normalise(text):\n    text = text.lower()\n    text = re.sub(r\'http\\S+|www\\S+\', \'\', text)   # remove URLs\n    text = re.sub(r\'@\\w+\', \'\', text)               # remove @mentions\n    text = re.sub(r\'[^a-z0-9\\s]\', \'\', text)       # alphanumeric only\n    text = re.sub(r\'\\d+\', \'\', text)                # remove numerics\n    tokens = text.split()\n    tokens = [w for w in tokens if len(w) >= 3]    # min length filter\n    return tokens')

h2('Chunking — Noun Phrase Extraction')
code('import nltk\nfrom nltk import RegexpParser, pos_tag, word_tokenize\n\n# Grammar: optional DT + any JJ + noun\ngrammar = "NP: {<DT>?<JJ>*<NN>}"\nparser  = RegexpParser(grammar)\n\ntokens = word_tokenize("The quick brown fox jumped over the lazy dog")\ntagged = pos_tag(tokens)    # [(\'The\',\'DT\'),(\'quick\',\'JJ\'),(\'fox\',\'NN\'),...]\ntree   = parser.parse(tagged)\n\n# Extract noun phrase chunks\nfor subtree in tree.subtrees():\n    if subtree.label() == \'NP\':\n        print(\' \'.join(w for w,t in subtree.leaves()))')

h2('POS-Tag Filtering (Keep Content Words Only)')
code('KEEP_TAGS = {\'NN\',\'NNS\',\'NNP\',\'NNPS\',          # nouns\n             \'VB\',\'VBD\',\'VBG\',\'VBN\',\'VBP\',\'VBZ\', # verbs\n             \'JJ\',\'JJR\',\'JJS\'}                   # adjectives\n\ndef pos_filter(text):\n    tokens = word_tokenize(text.lower())\n    tagged = pos_tag(tokens)\n    return [word for word, tag in tagged if tag in KEEP_TAGS]')

h2('Text Preprocessing Pipeline')
for i, s in enumerate([
    'Data Cleaning: remove HTML, punctuation, numbers, special characters',
    'Tokenization: split text into tokens (words, subwords, or chars)',
    'Stop Word Removal: remove high-frequency, low-info words (a, an, the, is...)',
    'Stemming / Lemmatization: reduce words to root form',
    'POS Tagging: label each token (noun, verb, adjective...)',
    'Feature Extraction: convert text to numerical representation',
], 1):
    bullet(f'{i}. {s}')

callout('⚠ Stop words risk:', 'Blindly removing stop words destroys "President of University of Hong Kong", "Let it be", "flights to London". Use Named Entity Recognition (NER) to protect phrases first.', 'FFFBEB')

h2('Stemming vs Lemmatization')
table(
    ['', 'Stemming', 'Lemmatization'],
    [
        ['Method',   'Chop suffixes by rules',          'Dictionary lookup for root (lemma)'],
        ['Result',   'May NOT be a real word',           'Always a valid dictionary word'],
        ['Speed',    'Faster',                           'Slower (needs POS context)'],
        ['Example',  '"computational" → "comput"',       '"am/are/is" → "be"'],
        ['Algorithm','Porter Stemmer (5 phases)',        'WordNet Lemmatizer'],
    ],
    [3, 6, 6.2]
)

h2('Porter\'s Stemmer — 5 Phases')
table(
    ['Phase', 'Rules / Examples'],
    [
        ['1', 's → ∅ (cats→cat), sses→ss (classes→class), ies→i (companies→compani)'],
        ['2', 'ational→ate (relational→relate), tional→tion (conditional→condition)'],
        ['3', 'icate→ic (triplicate→triplic), alize→al (formalize→formal)'],
        ['4', '(m>1)ement→∅ (replacement→replac), ment→∅, ent→∅, ion→∅ (if preceded by s/t)'],
        ['5', 'Remove trailing e (probate→probat), double consonant cleanup'],
    ],
    [1.5, 13.7]
)

h2('Levenshtein (Edit) Distance')
bullet('Minimum single-character edits (insert, delete, substitute) to transform string s → t')
code('Recurrence:\n  if s[i]==t[j]:  dp[i][j] = dp[i-1][j-1]          (match, no cost)\n  else:           dp[i][j] = 1 + min(\n                      dp[i-1][j],    # delete from s\n                      dp[i][j-1],    # insert into s\n                      dp[i-1][j-1]   # substitute\n                  )\nExample: "kitten" → "sitting" = 3 edits (k→s, e→i, insert g)')

h2('Byte Pair Encoding (BPE) — Three-Phase Algorithm')
table(
    ['Component', 'Role', 'I/O'],
    [
        ['Token Learner',          'Processes corpus; finds most frequent adjacent pairs; builds merge rules + vocab', 'Raw text → merge rules'],
        ['Token Segmenter/Encoder','Applies learned merge rules to encode new sentences',                             'New text → token IDs'],
        ['Token Merger/Decoder',   'Reverses encoding back to readable text',                                         'Token IDs → text'],
    ],
    [4.5, 7, 3.7]
)
for i, s in enumerate([
    'Split all words into characters + Ġ (end-of-word boundary marker)',
    'Count all adjacent symbol pairs in corpus',
    'Merge the most frequent pair into a new symbol; add to vocabulary',
    'Repeat until desired vocabulary size',
], 1):
    bullet(f'{i}. {s}')

h3('BPE Worked Example')
code('Corpus: hug(x10), pug(x5), pun(x12), bun(x4), hugs(x5)\n\nInitial tokens: [Ġ,h,u,g] [Ġ,p,u,g] [Ġ,p,u,n] [Ġ,b,u,n] [Ġ,h,u,g,s]\n\nMost frequent pair: (u,g) = 20 times → Merge: ug\nNext: (Ġ,h) = 15 times → Merge: Ġh\nNext: (Ġh,ug) = 15 times → Merge: Ġhug\n\nFinal: "hug" → token_id=7   "pun" → [Ġp,u,n] = [9,3,5]\nOOV never happens — unknown words split into known subword tokens')

h2('Feature Representations')
table(
    ['Method', 'Description', 'Limitation'],
    [
        ['One-Hot Encoding', 'Word → binary vector of vocab size; only its position=1', 'Sparse; no semantic similarity'],
        ['Bag of Words',     'Count occurrences of each word; stop words removed',      'Word ORDER is lost; sparse'],
        ['TF-IDF',           'TF × IDF — high score = frequent in ONE doc, rare across docs', 'Still no word order'],
    ],
    [3.5, 7, 4.7]
)

h2('TF-IDF — Full Formula & Worked Example')
code('IDF(t,D)       = log( (|D|+1) / (DF(t,D)+1) )    ← +1 = Laplace smoothing\nTF-IDF(t,d,D)  = TF(t,d) x IDF(t,D)\n|D|=total docs, DF(t,D)=docs containing t\n\nExample (|D|=2):  D0="Python python Spark Spark"  D1="Python SQL"\n  TF(python,D0)=2 TF(spark,D0)=2 TF(python,D1)=1 TF(sql,D1)=1\n  DF(python)=2    DF(spark)=1    DF(sql)=1\n  IDF(python)=log(3/3)=0       <- in both docs, not discriminating\n  IDF(spark) =log(3/2)=0.405\n  IDF(sql)   =log(3/2)=0.405\n  TF-IDF(spark,D0)=2x0.405=0.811  <- "spark" is keyword for D0\n  TF-IDF(sql,D1) =1x0.405=0.405   <- "sql"   is keyword for D1')

h2('Spark ML NLP Pipeline')
table(
    ['Component', 'What it does', 'Key params'],
    [
        ['Tokenizer',          'Split sentence into word tokens',        'inputCol, outputCol'],
        ['StopWordsRemover',   'Remove high-freq, low-info words',       'inputCol, outputCol'],
        ['NGram(n=2)',         'Generate bigrams / n-word phrases',      'n=2 for bigrams'],
        ['CountVectorizer',    'Term frequency sparse vectors',           'minTF, minDF, vocabSize'],
        ['IDF',                'Apply inverse-doc-freq weighting',       'inputCol="tf"'],
    ],
    [4, 5.5, 5.7]
)
code('from pyspark.ml import Pipeline\nfrom pyspark.ml.feature import Tokenizer, StopWordsRemover, NGram, CountVectorizer, IDF\n\npipeline = Pipeline(stages=[\n    Tokenizer(inputCol="sentence", outputCol="words"),\n    StopWordsRemover(inputCol="words", outputCol="filtered"),\n    NGram(n=2, inputCol="filtered", outputCol="ngrams"),\n    CountVectorizer(minTF=1.0, minDF=1.0, vocabSize=20,\n                    inputCol="ngrams", outputCol="tf"),\n    IDF(inputCol="tf", outputCol="tfidf")\n])\nmodel = pipeline.fit(data)\nresult = model.transform(data)')

h2('UDFs (User-Defined Functions)')
code('from pyspark.sql.functions import udf\nfrom pyspark.sql.types import StringType\n\nmy_udf = udf(lambda text: text.upper(), StringType())\ndf.withColumn("upper_text", my_udf(df["text"])).show()')

h2('countByKey() vs countByValue()')
code('countByKey()   -> count occurrences of each KEY   {"Mazda":2, "Ferrari":1}\ncountByValue() -> count occurrences of each (key,value) PAIR\n\n# Document Frequency pattern:\ntokenized.flatMapValues(lambda x: x)   # one (docID,word) per word\n         .distinct()                    # unique (docID,word) pairs\n         .map(lambda x: (x[1],x[0]))   # flip to (word,docID)\n         .countByKey()                  # how many docs contain each word')

# ─────────────────────────────────────────────────────────────────────────
pg()
h1('L6 · LLMs & Word Embeddings (Word2Vec)')

h2('Distributional Hypothesis')
callout('"A word is characterized by the company it keeps."',
        'Firth (1957); Harris (1954): words in similar contexts have similar meanings. Underpins Word2Vec, GloVe, and all distributional models.', 'DBEAFE')

h2('Word Vector Analogy Arithmetic')
code('vec("king") - vec("man") + vec("woman") ~ vec("queen")\nvec("Paris") - vec("France") + vec("Italy") ~ vec("Rome")\n\nEmbeddings encode relational structure:\n  Royalty offset = king - man ~ queen - woman\n  Capital-of offset = Paris - France ~ Rome - Italy')
callout('One-hot limitation:', 'cosine similarity between ANY two distinct one-hot vectors = 0 (orthogonal). Embeddings place semantically similar words nearby in vector space.', 'FFFBEB')

h2('Special Tokens')
table(
    ['Token', 'Purpose'],
    [
        ['<PAD>', 'Pad shorter sequences to equal length for batching'],
        ['<UNK>', 'Replace rare / out-of-vocabulary words'],
        ['<BOS>/<EOS>', 'Mark beginning / end of a sequence'],
    ],
    [3.5, 11.7]
)

h2('GloVe — Global Vectors (Stanford)')
table(
    ['', 'Word2Vec', 'GloVe'],
    [
        ['Training signal', 'Local context window',      'Global co-occurrence matrix X'],
        ['Method',          'Predictive neural, SGD',    'Matrix factorization (least squares)'],
        ['Key advantage',   'Efficient, streams data',   'Captures full corpus statistics'],
    ],
    [3.5, 5.8, 5.9]
)
code('GloVe objective:\n  J = Σ_{i,j} f(X_ij)(w_i^T w_j + b_i + b_j − log X_ij)^2\n\nX_ij = co-occurrence count of word j in context of word i\nf(x)  = weighting function (caps frequent pairs)')

h2('CBOW Neural Architecture')
code('Layer 1: Embedding lookup  → one vector per context word (dim=d)\nLayer 2: Lambda mean pool   → x = mean(v_{w-m}, ..., v_{w+m})\nLayer 3: Dense (softmax)    → P(center | context) = softmax(W·x + b)\n\n- Input: 2m one-hot context vectors (m = window size)\n- Shared embedding matrix across all context words\n- Mean pooling loses word order ("bag" of words)')

h2('Text Similarity Algorithms')
table(
    ['Algorithm', 'Formula', 'Key constraint'],
    [
        ['Hamming Distance',   'Count positions where characters differ',              'Strings must be SAME length; substitutions only (no insert/delete)'],
        ['Sorensen-Dice',      '2|A∩B| / (|A|+|B|)',                                  'Uses bigram sets; range [0,1]; symmetric'],
        ['Tversky Index',      '|A∩B| / (|A∩B| + α|A\\B| + β|B\\A|)',               'Asymmetric; α=β=0.5→Dice, α=β=1→Jaccard'],
        ['Overlap Coefficient','|A∩B| / min(|A|,|B|)',                                '=1.0 if one set is subset of the other'],
    ],
    [3.8, 4.8, 6.6]
)
body('Levenshtein: character sequences. Hamming: equal-length strings only. Dice/Tversky/Overlap: token/bigram sets — fuzzy document matching.')

h2('Skip-gram vs CBOW')
table(
    ['', 'Skip-gram', 'CBOW'],
    [
        ['Input',       'Center word',                      'Context words (averaged)'],
        ['Output',      'Predict surrounding context words','Predict center word'],
        ['Better for',  'Rare words',                       'Frequent words'],
        ['Speed',       'Slower to train',                  'Faster to train'],
        ['Embeddings',  'Input vectors v_c preferred',      'Context vectors V_o used as embeddings'],
    ],
    [3, 6, 6.2]
)

h2('Skip-gram: Objective Function')
code('Maximize:  Σ_t  Σ_{-m≤j≤m, j≠0}  log P(o | c)\n\nP(o | c) = exp(u_o^T · v_c)  /  Σ_w exp(u_w^T · v_c)\n\n  v_c  = input embedding of center word c (from matrix V)\n  u_o  = output embedding of context word o (from matrix U)\n  Denominator = softmax over entire vocabulary (expensive for large vocab)')

h2('Skip-gram — Three Formal Assumptions')
bullet('1. i.i.d. windows: each context window is independent; objective sums log-probs across all windows')
bullet('2. Conditional independence: context words predicted independently given center word — P(o1,o2…|c) = ∏ P(oi|c)')
bullet('3. Position independence: word order / distance within window ignored — same prediction 1 or m steps away')
callout('Note:', 'These assumptions are linguistically wrong but make optimization tractable — and embeddings are still very good.', 'FEF3C7')

h2('CBOW — Numerical Example')
code('Vocab: {cat=0, sit=1, on=2, mats=3}   Center word = "on"\nContext: sit(1), mats(3)   window=1\n\nOne-hot: v_sit=[0,1,0,0]  v_mats=[0,0,0,1]\nEmbedding lookup: e_sit=[0.2,-0.1,0.5]  e_mats=[-0.3,0.4,0.1]\nMean pooling: x = (e_sit+e_mats)/2 = [-0.05, 0.15, 0.3]\nOutput logits = U·x → [0.1, 0.8, 0.3, -0.2]\nSoftmax → P = [0.18, 0.36, 0.22, 0.13]\n               cat   sit   on   mats\nPredicted: "sit" (wrong) → compute loss → backprop → update W, U')

h2('CBOW: Loss Function')
code('Input:  x = (1/2m) Σ v_w  for w in context window of size m\n\nP(c | context) = exp(u_c^T · x)  /  Σ_w exp(u_w^T · x)\n\nLoss: minimize  NLL = -log P(c | context)\n\nNote: context vectors V_o are the final word embeddings in CBOW')

h2('Skip-gram with Negative Sampling (SGNS)')
code('y = 1  for real (center, context) pairs\ny = 0  for k randomly sampled noise pairs\n\nLoss = -[ log σ(u_o^T v_c)  +  Σ_k log σ(-u_k^T v_c) ]\n\nσ(x) = 1 / (1 + e^(-x))    (sigmoid)\n\nWhy: converts expensive softmax (50k-class) into cheap binary classifiers')

h2('Word2Vec → Transformers (conceptual)')
bullet('Word2Vec: static embeddings — one vector per word regardless of context')
bullet('BERT/GPT: contextual embeddings — same word has different vector per sentence')
bullet('Attention: model looks at all positions simultaneously (no fixed window)')

# ─────────────────────────────────────────────────────────────────────────
h1('L7 · Link Analysis — PageRank')

h2('Web as a Graph & Links as Votes')
bullet('Nodes = web pages;  Edges = hyperlinks (directed)')
bullet('Links as Votes: a link from page A to page B = a vote by A for B')
bullet('Weighted votes: a link from an important page is worth more → recursive definition → PageRank')
bullet('More in-links = more important; link from important source = stronger vote')

h2('Random Walk Interpretation')
bullet('Imagine a random surfer: with prob β follow a random outgoing link; with prob (1−β) teleport to any page uniformly')
bullet('PageRank(page) = long-run fraction of time the surfer spends on that page')
bullet('This is the stationary distribution π of the Markov chain defined by the Google Matrix A')
code('Stationary distribution: π = A · π\nPower iteration finds it: r converges to π regardless of start\n(Ergodicity guaranteed by teleport term — every state reachable from every other state)')

h2('Core Idea')
body('A page is important if important pages link to it. PageRank distributes "rank" along edges until convergence.')

h2('Stochastic Adjacency Matrix M')
code('M[j][i] = 1 / out_degree(i)   if edge i → j exists\nM[j][i] = 0                   otherwise\n\nColumns sum to 1 (column-stochastic).\nPageRank vector r satisfies:  r = M · r  (eigenvector with eigenvalue 1)')

h2('Problems & Fixes')
table(
    ['Problem', 'Cause', 'Fix'],
    [
        ['Dead Ends',     'Node with no outgoing edges — absorbs rank',          'Treat as linking to all N pages uniformly (teleport)'],
        ['Spider Traps',  'Group linking only to each other — traps all rank',   'Random teleport with probability (1−β)'],
    ],
    [2.8, 7, 5.4]
)

h3('Dead-End Rank Redistribution (detail)')
bullet('Dead ends cause rank leakage — their accumulated rank is never passed on')
bullet('Per-iteration fix: leaked = (1 - sum(r_new)) / N; add this back to every page')
bullet('Google Matrix (1-β)/N term achieves same effect globally for all nodes')

h2('PageRank at Scale')
code('Web scale (1998):\n  ~1 billion pages\n  Rank vector: one float64 per page → ~8 GB for one copy of r\n  Two copies per iteration (r_old + r_new) → ~16 GB minimum\n\nWhy distributed:\n  Single machine cannot hold rank vector in RAM (1998)\n  Spark/Hadoop: partition adjacency list + rank vector across nodes\n  Only rank contributions (src→dst pairs) cross nodes each iteration')

h2('Google Matrix')
code('A = β · M  +  (1−β) · [1/N]_{N×N}\nr = A · r\n\nβ ∈ [0.8, 0.9]  — damping factor\n(1−β)           — teleport probability\n\nInterpretation: prob β → follow a link; prob (1−β) → jump to random page')

h2('Power Iteration (Algorithm)')
for i, s in enumerate([
    'Initialise: r = [1/N, 1/N, …, 1/N]',
    'Multiply: r_new = A · r',
    'Repeat until ||r_new − r||₁ < ε (converges ~50 iterations on web graph)',
], 1):
    bullet(f'{i}. {s}')

h2('Efficient Sparse Formula (Tutorial 7)')
code('Avoids dense N×N matrix:\n  r^(t+1) = β · M · r^t  +  (1−β)/N\n\nPySpark implementation:\n  contribs = adj_list.join(ranks).flatMap(\n      lambda (url, (links, rank)): [(link, rank/len(links)) for link in links]\n  )\n  ranks = contribs.reduceByKey(add).mapValues(\n      lambda r: beta*r + (1-beta)/N\n  )')

h2('Spider Trap — Numerical Example')
code('Graph: A→B, B→A, B→m, m→m   (m links only to itself)\nStart: r = [1/3, 1/3, 1/3]\n\nIter 1: r_A=1/6, r_B=1/3, r_m=1/2\nIter 2: r_A=1/6, r_B=1/6, r_m=2/3\nConverges to: r = [0, 0, 1]   ALL rank absorbed by trap!\n\nFix: teleport with β=0.85 — surfer occasionally escapes')

h2('Dead End — Numerical Example')
code('Graph: A→B, B→C, C→∅   (C has no outgoing links)\nStart: r = [1/3, 1/3, 1/3]\n\nIter 1: r = [0, 1/3, 1/3]   sum=2/3 — rank leaked!\nIter 2: r = [0, 0, 1/3]\nConverges to: r = [0, 0, 0]   rank disappears entirely!\n\nFix: treat dead-end as teleporting to all N pages uniformly')

h2('Convergence — Numerical Trace (β=0.8)')
code('Graph: A→B, A→C, B→C, C→A\nGoogle Matrix A = 0.8·M + 0.2·[1/3]\n\nIter 0: r = [0.333, 0.333, 0.333]\nIter 1: r = [0.333, 0.200, 0.467]\nIter 2: r = [0.440, 0.194, 0.367]\nIter 3: r = [0.360, 0.230, 0.410]\n...\nConverged: r ≈ [7/33, 5/33, 21/33] = [0.212, 0.152, 0.636]\nC has highest rank — receives links from both A and B')

h2('Worked Example (3-node graph)')
code('Nodes: A, B, C\nEdges: A→B, A→C, B→C, C→A\n\n      A    B    C\nM = [ 0    0    1  ]   (who points TO A)\n    [1/2   0    0  ]   (who points TO B)\n    [1/2   1    0  ]   (who points TO C)\n\nStart: r = [1/3, 1/3, 1/3]^T\nIter 1: M·r = [1/3, 1/6, 1/2]^T  → apply Google Matrix with β=0.85\n...(iterate until convergence)')

callout('★ Exam question type:', 'Given a small graph, build M by hand, run 1–2 power iterations, identify dead ends / spider traps.', 'DBEAFE')

# ─────────────────────────────────────────────────────────────────────────
pg()
h1('L8 · Graph Analytics')

h2('Graph Types')
table(
    ['Type', 'Description', 'Degree concept', 'Examples'],
    [
        ['Undirected', 'Edges have no direction; A–B = B–A',  'Single degree d(v)',                      'Friendship, co-authorship'],
        ['Directed',   'Edges have direction; A→B ≠ B→A',    'In-degree (edges TO v) + Out-degree (FROM v)', 'Twitter follow, hyperlinks, citations'],
        ['Unweighted', 'All edges have equal weight (1)',      'Distance = hop count',                    'Basic social graphs'],
        ['Weighted',   'Edges carry numeric weights',          'Use Dijkstra for shortest weighted paths', 'Road networks, airline routes'],
    ],
    [2.5, 5, 3.5, 4.2]
)
code('For directed graphs:\n  Σ in-degree(v)  =  Σ out-degree(v)  =  |E|\n  (each edge contributes 1 to in-degree of target AND 1 to out-degree of source)')

h2('Graph Basics')
table(
    ['Term', 'Definition', 'Formula'],
    [
        ['Degree d(v)',       'Number of edges incident to v',                  '—'],
        ['Average Degree',    'Mean degree across all nodes',                   '2|E| / N'],
        ['Graph Density',     'Fraction of possible edges that exist',          '2|E| / (N(N−1))'],
        ['Distance d(u,v)',   'Edges on shortest path between u and v',         'BFS'],
        ['Diameter',          'Maximum shortest-path distance in graph',        'max over all pairs d(u,v)'],
        ['Avg Path Length',   'Average distance between all node pairs',        'mean of all d(u,v)'],
    ],
    [3.5, 6.5, 5.2]
)

h2('Graph Connectivity — Weakly vs Strongly Connected')
table(
    ['Type', 'Definition', 'Example'],
    [
        ['Weakly connected',    'Connected if edge directions are ignored (undirected path exists)',    'Most real directed graphs (web, Twitter)'],
        ['Strongly connected',  'Directed path from EVERY node to EVERY other node',                   'Stricter; web graph is not strongly connected'],
    ],
    [3.5, 7, 4.7]
)

h2('Clustering Coefficient')
code('Local (per-node):\n  C(v) = 2·m_v / (d_v · (d_v - 1))\n  m_v = edges between v\'s neighbors\n  d_v = degree of v  (need d_v >= 2)\n  Range [0,1]:  1 = all neighbors connected; 0 = none\n\nGlobal (graph-wide):\n  C_global = (1/N) · Σ_v C(v)   <- average of all local coefficients\n  Measures overall "cliquishness" of the network\n  Small-world networks: high C_global + small average path length')

h2('Centrality Measures')
table(
    ['Measure', 'Formula', 'Limitation'],
    [
        ['Closeness',    'C_clos(v) = (N−1) / Σ_x d(v,x)',              'Undefined for disconnected graphs'],
        ['Harmonic',     'C_harm(v) = Σ_{x≠v} 1/d(v,x)',               'Handles disconnected (unreachable → 0)'],
        ['Betweenness',  'BC(v) = Σ_{s≠t} σ_st(v) / σ_st',            'O(VE) to compute; expensive'],
        ['Eigenvector',  'x_v = (1/λ) Σ_{u∈N(v)} x_u',               'Ill-defined for disconnected graphs'],
    ],
    [3, 6.5, 5.7]
)
body('σ_st = total shortest paths from s to t; σ_st(v) = those passing through v')

h2('Eigenvector Centrality — Power Iteration')
code('c^(0) = [1, 1, ..., 1]^T   (initialise uniformly)\nc^(t+1) = A · c^(t)  /  || A · c^(t) ||   (multiply then normalise)\n\nRepeat until convergence.\nResult = eigenvector for largest eigenvalue of adjacency matrix A.\n\nA[i][j]=1 if edge exists; captures global influence (PageRank is a variant)')
bullet('Low-degree node can have high eigenvector centrality if connected to hubs')

h2('Edge Betweenness & Community Detection')
code('EB(u,v) = Σ_{s≠t}  σ_st(u,v) / σ_st\n\nAlgorithm (BFS fractional counting):\n  1. Run BFS from each source s\n  2. For each target t: credit = 1 / σ_st\n  3. Distribute credit backwards along BFS tree edges\n  4. Edges on more shortest paths accumulate higher EB\n\nGirvan-Newman community detection:\n  Repeatedly remove highest-EB edge → graph splits into communities')

h2('Real-World Case Study: MSN Messenger (2006)')
table(
    ['Property', 'Value'],
    [
        ['Nodes (users)',                '~180 million'],
        ['Edges (communication pairs)',  '~1.3 billion'],
        ['Average path length',          '6.6 hops'],
        ['% within 8 hops',             '~90%'],
    ],
    [5.5, 9.7]
)
callout('Six Degrees of Separation:', 'Confirms small-world property at massive scale. Most pairs of people connected through ~6 intermediaries.', 'D1FAE5')

h2('Centrality Limitations — Krackhardt\'s Kite Graph')
table(
    ['Centrality', 'Identifies', 'Highest-scoring node role'],
    [
        ['Degree',       'Most directly connected',          'Hub node (many neighbours)'],
        ['Closeness',    'Fastest info spreader',            'Central node (short paths to all)'],
        ['Betweenness',  'Controls information flow',        'Bridge / broker between clusters'],
        ['Eigenvector',  'Most globally influential',        'Node connected to other influencers'],
    ],
    [3.5, 5, 6.7]
)
callout('Key insight:', 'No single centrality is "best" — four measures give four different rankings in the Krackhardt kite graph. Choose based on what "important" means.', 'FFFBEB')

h2('Betweenness — Worked Example')
body('Key insight: high betweenness = "bridge" node. Removing it fragments the network.')
code('Linear chain: A–B–C–D–E with shortcut B–D\nBC(B) sums contributions from all (s,t) pairs where s≠B≠t:\n  (A,C): only path A→B→C; σ(B)=1/1 = 1.0\n  (A,D): shortest A→B→D; σ(B)=1/1 = 1.0\n  (A,E): only path A→B→D→E; σ(B)=1/1 = 1.0\n  (C,E): C→D→E; does not pass B; σ(B)=0\n  Add all pairs → BC(B)')

h2('Graph Connectivity')
bullet('Connected graph: path exists between every pair of nodes')
bullet('Giant component: largest connected subgraph (contains most nodes in social/web graphs)')
bullet('BFS (Breadth-First Search) used to find connected components and shortest paths')
code('BFS from s:\n  1. Mark s visited, enqueue\n  2. While queue not empty:\n       u = dequeue\n       for each neighbour v of u not yet visited:\n           mark visited; distance[v] = distance[u]+1; enqueue v')

h2('Degree Distribution & Network Types')
table(
    ['Network type', 'Degree distribution', 'Properties', 'Examples'],
    [
        ['Random',      'Bell curve (Poisson)',   'Uniform connectivity',                  'Erdős–Rényi model'],
        ['Small-World', 'Bell curve',             'Short paths + high clustering coefficient','Social networks'],
        ['Scale-Free',  'Power law P(k)~k^(−γ)', 'Few hubs with very high degree; heavy tail','WWW, citations'],
    ],
    [3, 4, 5, 3.2]
)

# ─────────────────────────────────────────────────────────────────────────
h1('L9 · Big Data Storage & NoSQL')

h2('ACID Properties')
table(
    ['Property', 'Meaning'],
    [
        ['Atomicity',    'All operations succeed or none (all-or-nothing)'],
        ['Consistency',  'DB moves between valid states; all rules enforced'],
        ['Isolation',    'Concurrent transactions execute as if sequential'],
        ['Durability',   'Committed transactions survive system failure'],
    ],
    [3.5, 11.7]
)

h2('CAP Theorem')
body('In a distributed system you can guarantee AT MOST 2 of 3:')
table(
    ['Letter', 'Meaning'],
    [
        ['C — Consistency',          'Every read returns most recent write (or error)'],
        ['A — Availability',         'Every request gets a non-error response (may be stale)'],
        ['P — Partition Tolerance',  'System operates despite dropped/delayed messages between nodes'],
    ],
    [4, 11.2]
)
table(
    ['Choose', 'Give Up', 'Typical System'],
    [
        ['C + P', 'Availability (during partition)',  'MySQL, PostgreSQL, HBase, Zookeeper → RDBMS'],
        ['A + P', 'Consistency (may return stale)',   'Cassandra, DynamoDB, CouchDB → NoSQL'],
        ['C + A', 'Partition tolerance (impossible)', 'Single-machine RDBMS only'],
    ],
    [2, 5, 8.2]
)
callout('Key rule:', 'Network partitions WILL happen. Real distributed systems choose C vs A. CA = theoretical only.', 'FEF2F2')

h2('Strong vs Eventual Consistency')
table(
    ['', 'Strong Consistency', 'Eventual Consistency'],
    [
        ['Guarantee', 'Read after write always returns new value; all nodes in sync', 'Converges over time; reads may be temporarily stale'],
        ['Latency',   'Higher; may block writes',                                     'Lower; always writable'],
        ['Used by',   'RDBMS (CP systems)',                                           'NoSQL (AP systems)'],
    ],
    [3, 6, 6.2]
)

h2('BASE Model (NoSQL alternative to ACID)')
table(
    ['Letter', 'Stands for', 'Meaning'],
    [
        ['B', 'Basically Available', 'System always acknowledges requests (may return stale data)'],
        ['S', 'Soft State',          'State can change over time even without new input (propagation)'],
        ['E', 'Eventually Consistent','Reaches consistency once all propagation completes'],
    ],
    [1.5, 3.5, 10.2]
)

h2('NoSQL Data Models')
table(
    ['Type', 'Structure', 'Examples', 'Use case'],
    [
        ['Key-Value',    '{key: value}',          'Redis, DynamoDB',    'Sessions, caching'],
        ['Document',     'JSON/BSON documents',   'MongoDB, CouchDB',   'Content, catalogs'],
        ['Column-Family','Rows + column families', 'Cassandra, HBase',  'Time-series, analytics'],
        ['Graph',        'Nodes + edges',          'Neo4j',             'Social networks, recommendations'],
    ],
    [3, 4, 4, 4.2]
)

h2('Partitioning vs Replication Trade-offs')
table(
    ['Strategy', 'Improves', 'Hurts'],
    [
        ['Partitioning (sharding)', 'Write performance',                     'Reads that join across partitions (expensive)'],
        ['Replication',             'Fault tolerance + read performance',    'Write performance (must update all replicas)'],
    ],
    [4, 5, 6.2]
)

h2('Transactions & OLTP vs OLAP')
code('BEGIN TRANSACTION\n  UPDATE accounts SET balance = balance - 100 WHERE id = "A";\n  UPDATE accounts SET balance = balance + 100 WHERE id = "B";\nCOMMIT      <- makes permanent\n-- or ABORT/ROLLBACK  <- undoes all changes on failure')
table(
    ['', 'OLTP', 'OLAP'],
    [
        ['Workload',  'Many short read/write transactions (ms)',  'Few complex read-heavy queries over large data'],
        ['Examples',  'Banking, e-commerce, booking systems',     'Data warehouses, reporting, Spark/Hive'],
        ['ACID',      'Requires full ACID',                       'Relaxed (eventual consistency acceptable)'],
    ],
    [3, 6, 6.2]
)

h2('Key-Value Store Operations')
table(
    ['Op', 'Description'],
    [
        ['Get(key)',            'Retrieve value for a single key'],
        ['Put(key, value)',     'Insert or update a key-value pair'],
        ['Delete(key)',        'Remove a key-value pair'],
        ['Multi-get([keys])',  'Batch retrieve multiple values in one call'],
        ['Multi-put({k:v})',   'Batch insert/update multiple pairs'],
        ['Range(k1, k2)',      'Retrieve all keys between k1 and k2 (sorted store only)'],
    ],
    [4.5, 10.7]
)

h2('3-Tiered Web Architecture')
code('Tier 1: Client (browser / mobile app)\n       ↕  HTTP\nTier 2: Application Server (business logic, stateless)\n       ↕  DB queries\nTier 3: Database (persistence layer)\n\nWhy it matters:\n  App servers are stateless → horizontally scalable\n  DB becomes bottleneck at scale → motivates NoSQL + caching\n  NoSQL removes strict ACID so Tier 3 also scales horizontally')

h2('XML vs JSON')
table(
    ['', 'XML', 'JSON'],
    [
        ['Verbosity', 'Every value needs open+close tags', 'Compact key-value pairs'],
        ['Parsing',   'Heavier, slower',                    'Native to JavaScript, lighter'],
        ['Usage',     'Legacy enterprise (SOAP, config)',    'REST APIs, MongoDB, Elasticsearch'],
        ['Trend',     'Being phased out',                    '"Many apps have replaced XML with JSON"'],
    ],
    [3.5, 5.5, 6.2]
)

# ─────────────────────────────────────────────────────────────────────────
pg()
h1('L10 · Recommendation Systems')

h2('Three Key Problems')
for i, s in enumerate([
    'Gathering ratings: collecting known user-item data (explicit asks vs implicit inference)',
    'Estimating unknowns: predicting missing ratings from known ones (core ML problem)',
    'Evaluation: measuring recommendation quality (RMSE, Precision@K, Coverage)',
], 1):
    bullet(f'{i}. {s}')

h2('Utility Matrix & Ratings')
body('Utility matrix: m users × n items; R[i,j] = rating. Most entries UNKNOWN (sparsity >99% on Netflix/Amazon).')
table(
    ['Rating type', 'Description', 'Examples', 'Problem'],
    [
        ['Explicit', 'User deliberately provides rating',     'Stars 1-5, thumbs up/down',        'Sparse — users rarely bother'],
        ['Implicit', 'Inferred from behaviour',               'Watch time, clicks, purchases',     'Cannot infer dislikes (not buying ≠ dislike)'],
    ],
    [2.8, 4, 3.5, 4.9]
)
callout('Implicit feedback asymmetry:', 'Can observe positive signals (bought/clicked). Not buying = dislike OR unseen OR unaffordable. Hard to learn from negatives.', 'FEF3C7')

h2('Two Main Approaches')
table(
    ['', 'Content-Based Filtering', 'Collaborative Filtering (CF)'],
    [
        ['Basis',        'Item features (genre, tags, director)',   'User behaviour (ratings matrix only)'],
        ['Data needed',  'Item metadata',                           'Other users\' ratings'],
        ['Cold start',   'Bad for new users; OK for new items',     'Bad for BOTH new users and new items'],
        ['Problem',      'Limited novelty (filter bubble)',         'Sparsity — most users rate very few items'],
    ],
    [3.5, 6, 5.7]
)

callout('CF Key Assumption:', '"If user A has same opinion as user B on item X, A is more likely to agree with B on a different item Y." — Similar past behaviour predicts similar future preferences.', 'DBEAFE')

h2('Collaborative Filtering — User-Based vs Item-Based')
table(
    ['', 'User-Based CF', 'Item-Based CF'],
    [
        ['Idea',        'Find similar users; recommend what they liked',     'Find similar items to those user rated'],
        ['Similarity',  'Between users (row vectors)',                       'Between items (column vectors)'],
        ['Stability',   'User prefs change → less stable',                  'Item characteristics stable → more stable'],
        ['Scale',       'Costly as users grow',                             'Pre-compute item-item matrix (Amazon\'s approach)'],
    ],
    [3, 6, 6.2]
)

h2('Similarity Metrics')
code('Jaccard Similarity (binary / implicit data):\n  sim(A,B) = |A intersect B| / |A union B|\n  Items rated by BOTH / Items rated by EITHER\n  Range [0,1]; ignores rating VALUES -> good for implicit\n\nCosine Similarity:\n  sim(u,v) = (u.v) / (||u|| * ||v||)\n  Range [-1,1]; 1=identical, 0=orthogonal, -1=opposite\n\nPearson Correlation (mean-centred cosine):\n  sim(u,v) = Sum(r_ui - r_u_bar)(r_vi - r_v_bar) / (sigma_u * sigma_v)\n  Adjusts for harsh vs generous raters')

h2('KNN-Based Collaborative Filtering')
code('User-based prediction for user u on item i:\n  r_hat_ui = r_bar_u  +  Σ_{v in kNN(u)} sim(u,v)*(r_vi - r_bar_v)\n                         ─────────────────────────────────────────\n                                 Σ_{v in kNN(u)} |sim(u,v)|\n\nr_bar_u = mean rating of user u (normalises harsh/generous raters)\nkNN(u)  = k most similar users (k typically 20-50)\n\nSteps:\n  1. Compute similarity between u and ALL other users\n  2. Pick top-k most similar users\n  3. Weighted average of their ratings for item i\n\nLimitation: O(N) similarity computation per prediction — does not scale to millions of users. ALS is far more scalable.')

h2('Matrix Factorization (ALS)')
code('Decompose ratings matrix R ≈ U · V^T\n  U = (users × k)  latent user matrix\n  V = (items × k)  latent item matrix\n  k = number of latent factors (50–200)\n\nPrediction: r̂_ui = u_i^T · v_j\n\nObjective: minimize Σ(r_ui − u_i^T v_j)² + λ(||U||²+||V||²)\n  optimised with ALS (Alternating Least Squares) or SGD')

h2('ALS Algorithm Steps')
for i, s in enumerate([
    'Initialise U and V with random values',
    'Fix V, solve for U (least squares on each user row)',
    'Fix U, solve for V (least squares on each item column)',
    'Repeat steps 2–3 until loss converges',
], 1):
    bullet(f'{i}. {s}')

h2('ALS in PySpark (Tutorial 9)')
code('from pyspark.ml.recommendation import ALS\nfrom pyspark.ml.evaluation import RegressionEvaluator\n\n(training, test) = data.randomSplit([0.8, 0.2], seed=1234)\nals = ALS(maxIter=5, regParam=0.01,\n          userCol="userId", itemCol="movieId", ratingCol="rating")\nmodel = als.fit(training)\npredictions = model.transform(test)\nrmse = RegressionEvaluator(metricName="rmse", labelCol="rating",\n    predictionCol="prediction").evaluate(predictions)\nprint("RMSE:", rmse)')

h2('Baseline Estimate')
code('b_ui = mu + b_u + b_i\n  mu  = global average rating\n  b_u = user bias (harsh/generous vs average)\n  b_i = item bias (popular/unpopular vs average)\n\nImproved CF prediction: r_hat_ui = b_ui + (CF adjustment from neighbours)')

h2('Evaluation Metrics')
table(
    ['Metric', 'Formula', 'Measures'],
    [
        ['RMSE',         'sqrt(Sum(r_hat-r)^2/n)',     'Prediction accuracy (lower=better)'],
        ['MAE',          'Sum|r_hat-r|/n',              'Mean absolute error'],
        ['Precision@K',  'Relevant in top-K / K',       'What fraction of top-K are actually good?'],
        ['Recall@K',     'Relevant in top-K / total',   'Did we surface all the good items?'],
        ['Coverage',     'Distinct items rec / total',  'Diversity — avoids always recommending popular items'],
    ],
    [3, 4, 8.2]
)

h2('Challenges')
table(
    ['Challenge', 'Description', 'Fix'],
    [
        ['Cold Start',     'New user/item has no ratings',           'Ask for initial prefs; use item content'],
        ['Sparsity',       'Ratings matrix >99% empty',              'Matrix factorization (ALS)'],
        ['Scalability',    'Millions of users × millions of items',  'Approximate nearest neighbours; distributed ALS'],
        ['Filter Bubble',  'Only recommends what\'s similar',        'Inject serendipitous recommendations'],
    ],
    [3, 5.5, 6.7]
)

# ─────────────────────────────────────────────────────────────────────────
pg()
h1('Tutorial Code Reference (T01–T04)')

h2('T01 — Pandas Essentials')
code('# Read TSV\ndf = pd.read_csv("data.tsv", delimiter="\\t")\n\n# Filter + sort\ndf[(df["gender"]=="F") & (df["year"]>1990)].sort_values("year", ascending=False)\n\n# Merge\nmerged = pd.merge(left=df, right=ratings,\n                  left_on=["userID","movieID"], right_on=["userID","movieID"])\n\n# Groupby mean\nmerged.groupby("name")[["rating"]].mean().sort_index()\n\n# Missing values\ndf.fillna("-")                 # replace NaN\ndf.dropna(axis=0)              # drop any-NaN rows\ndf.dropna(subset=["col"])      # drop where specific col is NaN\n\n# df["col"]   → Series (1D)\n# df[["col"]] → DataFrame (2D)')

h2('T02 — Dask')
code('import dask.array as da\narr = da.random.normal(20, 0.1, size=(20000,20000), chunks=(10000,10000))\narr.mean(axis=0).compute()\n\n# Dask DataFrame\nddf[~ddf["Cancelled"]].groupby("Origin").Origin.count().compute()\nddf.groupby("DayOfWeek").DepDelay.mean().idxmax().compute()\n\n# map_partitions (apply func to each pandas partition)\nmeta = pd.Series(name="Distance", dtype="float64")\nddf.Distance.map_partitions(lambda df: df*1.60934, meta=meta)')

h2('T03 — Spark RDD (Word Count, DNA, Palindrome)')
code('# Word count\ntextRDD.flatMap(lambda x: x.split())\\\n       .map(lambda x: (x.lower(),1))\\\n       .reduceByKey(add)\\\n       .sortBy(lambda x: x[1], ascending=False)\\\n       .collect()\n\n# DNA k-mer counting (k=5)\ndef generatePattern(line, k):\n    return [(line[i:i+k],1) for i in range(len(line)-k+1)]\ndnaRDD.flatMap(lambda l: generatePattern(l,5)).reduceByKey(add)\\\n      .sortBy(lambda x: x[1], ascending=False)\n\n# Palindrome filter\nseqRDD.filter(lambda x: x[1]>1)\\\n      .filter(lambda x: x[0]==x[0][::-1])\\\n      .collect()')

h2('T04 — Spark SQL / DataFrame')
code('from pyspark.sql import Row\nfrom pyspark.sql.functions import col, when\n\n# RDD → DataFrame\nrdd.map(lambda x: Row(sequence=x[0], count=x[1])).toDF()\n\n# Read CSV\ndf = ss.read.csv("file.txt", header=True, inferSchema=True)\n\n# Computed columns\ndf = df.withColumn("avg_time", (col("time")+col("timef"))/2)\ndf = df.withColumn("level",\n    when(col("score")<200,"Easy")\n    .when(col("score")<350,"Moderate")\n    .otherwise("Difficult"))\n\n# Aggregate + sort\ndf.groupby("name").avg("speed").sort("avg(speed)", ascending=False).show(5)\n\n# Union (same schema)\ncombined = df1.union(df2)\n\n# TF (term frequency per document)\ntf = tokenized.flatMapValues(lambda x: x).countByValue()\n# DF (document frequency — how many docs contain each term)\ndoc_freq = tokenized.flatMapValues(lambda x: x).distinct()\\\n                    .map(lambda x:(x[1],x[0])).countByKey()')

h2('Sparse Vector Operations (CountVectorizer output)')
code('# SparseVector format: (vocabSize, [indices], [values])\n# e.g. SparseVector(5, [0,2], [1.0,2.0]) = word0 once, word2 twice\nvec.toArray()     # [1.0, 0.0, 2.0, 0.0, 0.0]\nvec.indices       # [0, 2]\n\n# UDF to map index → term using vocabulary\ndef terms_idx2term(vocab):\n    def f(indices): return [vocab[i] for i in indices]\n    return udf(f, ArrayType(StringType()))\ndf.withColumn(\'terms\', terms_idx2term(cv_model.vocabulary)(df[\'indices\']))')

h2('UDF with Complex Return Type')
code('from pyspark.sql.types import ArrayType, StructType, StructField, StringType\nimport nltk\n\npos_schema = ArrayType(StructType([\n    StructField("word", StringType()),\n    StructField("tag",  StringType())\n]))\n\n@udf(pos_schema)\ndef pos_tag_udf(words):\n    return [{"word":w,"tag":t} for w,t in nltk.pos_tag(words)]\n\ndf.withColumn("pos_tags", pos_tag_udf(df["words"]))')

h2('Manual TF-IDF via RDD')
code('# Tokenize: (doc_id, [word1, word2, ...])\ntokenized = data.rdd.map(lambda r: (r.doc_id, r.text.lower().split()))\n\n# TF: count occurrences of each (doc, word) pair\ntf = tokenized.flatMapValues(lambda w: w).countByValue()\n# {(doc_id, word): count}\n\n# DF: how many docs contain each word\ndf_cnt = tokenized.flatMapValues(lambda w: set(w))\\\n                  .map(lambda x: (x[1],x[0])).countByKey()\n# {word: num_docs}\n\n# TF-IDF = TF * log((N+1)/(DF+1))\nimport math\nN = tokenized.count()\ntfidf = {(doc,w): cnt * math.log((N+1)/(df_cnt[w]+1))\n         for (doc,w),cnt in tf.items()}')

h2('Word2Vec — Nearest Words (Cosine Similarity)')
code('import numpy as np\n\n# Extract embedding matrix after training\nW = model.layers[0].get_weights()[0]  # (vocab_size, embed_dim)\n\n# L2-normalize for cosine similarity\nW_norm = W / (np.linalg.norm(W, axis=1, keepdims=True) + 1e-8)\n\n# Pre-compute similarity matrix\nsim_matrix = np.dot(W_norm, W_norm.T)  # (vocab_size, vocab_size)\n\ndef nearest_words(word, w2i, i2w, k=5):\n    idx = w2i[word]\n    sims = sim_matrix[idx]\n    top_k = np.argsort(sims)[::-1][1:k+1]\n    return [(i2w[i], sims[i]) for i in top_k]\n\ndef compare_words(w1, w2, w2i):\n    return float(np.dot(W_norm[w2i[w1]], W_norm[w2i[w2]]))\n\n# Analogy: king - man + woman ≈ queen\nv = W_norm[w2i[\'king\']] - W_norm[w2i[\'man\']] + W_norm[w2i[\'woman\']]\nprint(i2w[np.argmax(np.dot(W_norm, v))])')

h2('NetworkX — Graph Construction Exercises')
code('import networkx as nx\n\n# Star graph: hub at 0, spokes 1-7\ng8 = nx.Graph()\nfor i in range(1,8): g8.add_edge(0, i)\n\n# Circle (ring) graph\ng9 = nx.cycle_graph(9)\n\n# Complete graph (every pair connected)\ng10 = nx.complete_graph(10)\n\n# Linear chain\ng11 = nx.path_graph(11)\n\n# Graph metric functions\ndef avg_degree(g):   return 2*g.number_of_edges() / g.number_of_nodes()\ndef get_density(g):  N=g.number_of_nodes(); return 2*g.number_of_edges()/(N*(N-1))\n\n# Draw with centrality-scaled node sizes\ndef draw_network(G, centrality):\n    import matplotlib.pyplot as plt\n    pos = nx.spring_layout(G)\n    sizes = [v**2*10000 if v<1 else v*100 for v in centrality.values()]\n    nx.draw(G, pos, node_size=sizes, with_labels=True)\n    plt.show()\n\n# Average (global) clustering coefficient\nnx.average_clustering(g9)')

h2('Tutorial 6 — Keras Word Embedding Layer (CBOW)')
code('from tensorflow.keras.models import Sequential\nfrom tensorflow.keras.layers import Embedding, Lambda, Dense\nimport tensorflow.keras.backend as K\n\nvocab_size  = 5000   # vocabulary size\nembed_dim   = 100    # embedding dimension (typically 50-300)\nwindow_size = 2      # context window on each side\n\n# CBOW: context words -> mean pool -> predict center word\nmodel = Sequential([\n    Embedding(input_dim=vocab_size, output_dim=embed_dim,\n              input_length=2*window_size),   # 2m context words\n    Lambda(lambda x: K.mean(x, axis=1)),    # mean pooling\n    Dense(vocab_size, activation=\'softmax\') # output: prob over vocab\n])\nmodel.compile(optimizer=\'adam\', loss=\'sparse_categorical_crossentropy\')\n\n# After training: extract embedding weights\nembeddings = model.layers[0].get_weights()[0]  # (vocab_size, embed_dim)')

h2('Tutorial 7 — PageRank (PySpark + NetworkX)')
code('import networkx as nx\n\n# NetworkX built-in PageRank\nG  = nx.read_edgelist(\'links.txt\', create_using=nx.DiGraph())\npr = nx.pagerank(G, alpha=0.85, max_iter=100, tol=1e-6)\n# alpha=damping factor β; tol=convergence threshold\n\n# PySpark iterative with convergence check\nbeta, N = 0.85, num_nodes\nfor i in range(max_iter):\n    contribs = links.join(ranks).flatMap(\n        lambda u_ls_r: [(dest, u_ls_r[1][1]/len(u_ls_r[1][0]))\n                        for dest in u_ls_r[1][0]])\n    new_ranks = contribs.reduceByKey(lambda a,b: a+b)\\\n                        .mapValues(lambda r: beta*r + (1-beta)/N)\n    # Convergence: ||r_new - r_old||_1 < tol\n    delta = ranks.join(new_ranks).map(lambda x: abs(x[1][1]-x[1][0])).sum()\n    ranks = new_ranks\n    if delta < 1e-6:\n        break')

h2('Tutorial 8 — NetworkX Graph Analytics')
code('import networkx as nx\n\n# Build graph\nG = nx.Graph()                  # undirected; nx.DiGraph() for directed\nG.add_edge("Alice","Bob")\nG.add_edge("Bob","Carol")\nG.add_edges_from([("Dave","Eve"),("Eve","Frank")])\n\n# Basic properties\nG.number_of_nodes(), G.number_of_edges()\nnx.density(G)                   # 2|E| / (N(N-1))\n\n# Degrees\ndict(G.degree())                # {node: degree}\n\n# Centrality\ncc = nx.closeness_centrality(G)\nhc = nx.harmonic_centrality(G)  # handles disconnected graphs\nbc = nx.betweenness_centrality(G, normalized=True)\nec = nx.eigenvector_centrality(G, max_iter=100)\n\n# Clustering\nnx.clustering(G)                # per-node C(v)\nnx.average_clustering(G)\n\n# Shortest paths\nnx.shortest_path(G, "Alice", "Frank")\nnx.shortest_path_length(G, "Alice", "Frank")\n\n# Famous test graph\nK = nx.karate_club_graph()      # Zachary\'s karate club (34 nodes)\n\n# Girvan-Newman community detection\nfrom networkx.algorithms.community import girvan_newman\ncomp = girvan_newman(G)\ncommunities = tuple(next(comp))')

# ─────────────────────────────────────────────────────────────────────────
pg()
h1('★ Quick Reference — Formulas, Comparisons & Last-Minute Priority')

h2('Must-Know Formulas')
table(
    ['Topic', 'Formula'],
    [
        ['Average graph degree',         '2|E| / N'],
        ['Graph density',                '2|E| / (N(N−1))'],
        ['Clustering coefficient',       'C(v) = 2m_v / (d_v(d_v−1))'],
        ['Closeness centrality',         '(N−1) / Σ d(v,x)'],
        ['Harmonic centrality',          'Σ_{x≠v} 1/d(v,x)'],
        ['Betweenness centrality',       'Σ_{s≠t} σ_st(v) / σ_st'],
        ['PageRank (Google Matrix)',      'r = [βM + (1−β)(1/N)J] · r'],
        ['PageRank (sparse efficient)',   'r^(t+1) = β·M·r^t + (1−β)/N'],
        ['Skip-gram probability',        'exp(u_o^T v_c) / Σ_w exp(u_w^T v_c)'],
        ['SGNS loss',                    '−[log σ(u_o^T v_c) + Σ_k log σ(−u_k^T v_c)]'],
        ['Levenshtein recurrence',       'dp[i][j] = 1 + min(dp[i−1][j], dp[i][j−1], dp[i−1][j−1])'],
        ['Matrix factorization',         'R ≈ U·V^T;  r̂_ui = u_i^T · v_j'],
        ['TF-IDF',                       'TF(t,d) x log((|D|+1)/(DF(t,D)+1))'],
        ['Jaccard similarity',           '|A intersect B| / |A union B|'],
        ['Cosine similarity',            '(u·v) / (||u||·||v||)'],
    ],
    [5.5, 9.7]
)

h2('Key Comparisons')
table(
    ['A', 'vs', 'B', 'Key difference'],
    [
        ['Stemming',          'vs', 'Lemmatization',        'Stem may not be real word; lemma always is'],
        ['Skip-gram',         'vs', 'CBOW',                 'Skip: center→context; CBOW: context→center'],
        ['Narrow transform',  'vs', 'Wide transform',       'Wide requires network shuffle'],
        ['Transformation',    'vs', 'Action (Spark)',       'Transformation is lazy; action triggers execution'],
        ['ACID',              'vs', 'BASE',                 'ACID = strict; BASE = eventual consistency'],
        ['CP system',         'vs', 'AP system',            'CP sacrifices availability; AP sacrifices consistency'],
        ['Content-based',     'vs', 'Collaborative filter', 'Content: item features; CF: user behaviour'],
        ['Strong consistency','vs', 'Eventual consistency', 'Strong: always up-to-date; Eventual: converges over time'],
        ['RDD',               'vs', 'DataFrame',            'DF optimised by Catalyst; RDD is lower-level'],
        ['Data Lake',         'vs', 'Data Warehouse',       'Lake: raw schema-on-read; Warehouse: structured schema-on-write'],
        ['map()',             'vs', 'flatMap()',             'flatMap flattens one level of nested lists'],
        ['User-based CF',     'vs', 'Item-based CF',        'User prefs change; item characteristics are stable'],
    ],
    [3.8, 0.8, 3.8, 6.8]
)

h2('CAP Theorem Cheat Sheet')
table(
    ['Choose', 'Give Up', 'Typical System'],
    [
        ['C + A', 'P',  'Single-machine RDBMS (theoretical)'],
        ['C + P', 'A',  'MySQL, PostgreSQL, HBase, Zookeeper'],
        ['A + P', 'C',  'Cassandra, DynamoDB, CouchDB'],
    ],
    [2.5, 2.5, 10.2]
)

h2('Last-Minute Revision Priority')
for i, s in enumerate([
    '6 V\'s with examples (especially Veracity and Value)',
    'PageRank: Google Matrix formula, power iteration, dead ends vs spider traps',
    'CAP theorem — which systems are CP vs AP and why',
    'Centrality formulas: closeness, harmonic, betweenness (with example)',
    'Spark: all RDD ops, transformations vs actions, narrow vs wide',
    'Skip-gram vs CBOW (input/output, what embeddings are used)',
    'SGNS — why negative sampling, sigmoid function',
    'Levenshtein distance — DP table recurrence',
    'Collaborative vs content-based filtering + ALS steps',
    'ACID vs BASE; strong vs eventual consistency',
    'TF-IDF formula + worked example (IDF uses log with Laplace smoothing)',
    'Jaccard vs Cosine vs Pearson — when to use each',
    'Explicit vs implicit ratings; utility matrix sparsity',
    'Transactions: BEGIN/COMMIT/ABORT; OLTP vs OLAP',
    'BPE tokenization steps',
    'Porter\'s stemmer 5 phases with example rules',
    'MapReduce 5 steps + word count example',
    'HDFS: 3-way replication, write-once, fixed-size blocks',
    'Inverted index — flatMap + groupByKey pattern',
], 1):
    bullet(f'{i:2d}. {s}')

# ── Save ──────────────────────────────────────────────────────────────────
out = r'C:\Users\User\Downloads\4205\DSAI4205_Exam_Review.docx'
doc.save(out)
print(f'Saved to: {out}')
