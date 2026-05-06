# gen_mock_exam.py — DSAI4205 Mock Final Exam
# Format mirrors Dr. Fong's confirmed structure:
#   Section A: 22 MCQ (1 mark each)
#   Section B: Long questions (scenario, computation, programming)

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path('C:/Users/User/Downloads/4205/DSAI4205_Mock_Exam.docx')

# ── helpers ──────────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 90)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    run.font.size = Pt(7)

def q_heading(doc, num, topic, marks):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r1 = p.add_run(f'Question {num}  ')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    r2 = p.add_run(f'[{topic}]  ')
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r3 = p.add_run(f'({marks} marks)')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    r3.bold = True

def subq(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'({label})  ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)

def answer_box(doc, lines=4):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    cell.text = '\n' * lines
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
    cell_bg(cell, 'F8FAFC')
    doc.add_paragraph()

def ans_note(doc, text, color='green'):
    p = doc.add_paragraph()
    col = RGBColor(0x16, 0x65, 0x34) if color == 'green' else RGBColor(0x1E, 0x40, 0xAF)
    r1 = p.add_run('ANSWER: ')
    r1.bold = True
    r1.font.color.rgb = col
    r1.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.color.rgb = col
    r2.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)

def mcq(doc, num, stem, options, answer_letter, is_answer_key=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    r = p.add_run(f'{num}.  {stem}')
    r.font.size = Pt(10.5)
    for letter, text in options:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.8)
        p2.paragraph_format.space_after = Pt(1)
        is_correct = (letter == answer_letter)
        run = p2.add_run(f'{letter})  {text}')
        run.font.size = Pt(10)
        if is_answer_key and is_correct:
            run.bold = True
            run.font.color.rgb = RGBColor(0x16, 0x65, 0x34)

def dp_table(doc, rows_data, header_row=True):
    nrows = len(rows_data)
    ncols = len(rows_data[0])
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9.5)
            if ri == 0 or ci == 0:
                cell_bg(cell, 'DBEAFE')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()

# ── MCQ bank ─────────────────────────────────────────────────────────────────
MCQ_QUESTIONS = [
    # (stem, [(letter, option), ...], answer)
    # L1 — 2 questions
    (
        "Which of the following BEST describes the 'Veracity' dimension of Big Data?",
        [('A','The amount of data generated per second'),
         ('B','The diversity of data types including structured and unstructured'),
         ('C','The quality and accuracy of the data'),
         ('D','The usefulness of the data for business decision making')],
        'C'
    ),
    (
        "A company collects 2 petabytes of sensor data daily from 40,000 IoT devices. Which V does this PRIMARILY illustrate?",
        [('A','Velocity'), ('B','Variability'), ('C','Volume'), ('D','Veracity')],
        'C'
    ),
    # L2 — HDFS + Dask — 3 questions
    (
        "By default, HDFS stores each data block with how many replicas?",
        [('A','1'), ('B','2'), ('C','3'), ('D','5')],
        'C'
    ),
    (
        "Which of the following BEST describes Dask's execution model?",
        [('A','Eager — computations execute immediately when called'),
         ('B','Lazy — builds a task graph and only runs on .compute()'),
         ('C','Sequential — each operation completes before the next starts'),
         ('D','Synchronous — blocks the thread until all chunks are processed')],
        'B'
    ),
    (
        "Setting Dask chunk sizes too SMALL primarily causes which problem?",
        [('A','Out-of-memory errors from loading too much data'),
         ('B','Data corruption across chunk boundaries'),
         ('C','Excessive task scheduling overhead that slows computation'),
         ('D','Incorrect numerical results due to floating point errors')],
        'C'
    ),
    # L3 — MapReduce + RDD — 5 questions
    (
        "In MapReduce word count applied to the sentence 'big data big', what does the MAPPER emit?",
        [('A','[("big", 2), ("data", 1)]'),
         ('B','[("big", 1), ("data", 1), ("big", 1)]'),
         ('C','[("big data big", 1)]'),
         ('D','{"big": 2, "data": 1}')],
        'B'
    ),
    (
        "Which stage of MapReduce is AUTOMATICALLY handled by the framework with NO user code required?",
        [('A','Map'), ('B','Partition'), ('C','Shuffle and Sort'), ('D','Reduce')],
        'C'
    ),
    (
        "What is the PRIMARY reason MapReduce is slower than Spark?",
        [('A','MapReduce requires more CPU cores per task'),
         ('B','MapReduce writes all intermediate results to HDFS disk (with 3× replication)'),
         ('C','MapReduce cannot run on clusters larger than 100 nodes'),
         ('D','MapReduce loads the entire dataset into RAM before processing')],
        'B'
    ),
    (
        "Which of the following is a NARROW transformation in Spark RDD?",
        [('A','groupByKey()'), ('B','reduceByKey()'), ('C','filter()'), ('D','join()')],
        'C'
    ),
    (
        "What is the key difference between map() and mapValues() in Spark RDD?",
        [('A','map() is lazy; mapValues() executes immediately'),
         ('B','mapValues() applies the function only to the value, preserving the key; map() applies to the whole element'),
         ('C','map() works on any RDD; mapValues() only works on string RDDs'),
         ('D','mapValues() is faster because it skips serialization')],
        'B'
    ),
    # L4 — DataFrame — 2 questions
    (
        "The Catalyst optimizer in Spark SQL applies which of the following optimisations AUTOMATICALLY?",
        [('A','Predicate pushdown and column pruning'),
         ('B','Manual index creation on frequently-queried columns'),
         ('C','Converting Python UDFs to JVM bytecode'),
         ('D','Increasing the number of partitions before every join')],
        'A'
    ),
    (
        "Which statement about Spark DataFrame vs RDD is CORRECT?",
        [('A','RDD is always faster because it skips the serialisation layer'),
         ('B','DataFrame cannot handle nested data types such as arrays or structs'),
         ('C','DataFrame benefits from Catalyst optimisation; RDD does not'),
         ('D','Both RDD and DataFrame support SQL-style queries natively')],
        'C'
    ),
    # L5 — NLP — 3 questions
    (
        "In Byte Pair Encoding (BPE), what represents a subword or character that cannot be mapped to the learned vocabulary?",
        [('A','[MASK]'), ('B','[PAD]'), ('C','[UNK] / "?"'), ('D','[CLS]')],
        'C'
    ),
    (
        "What is the KEY difference between stemming and lemmatisation?",
        [('A','Stemming always produces a valid English word; lemmatisation may not'),
         ('B','Lemmatisation always produces a valid English word; stemming may not'),
         ('C','Stemming uses a dictionary; lemmatisation uses suffix-stripping rules'),
         ('D','Lemmatisation is faster but less accurate than stemming')],
        'B'
    ),
    (
        "In the Levenshtein distance dynamic programming table, where is the SOURCE word placed?",
        [('A','Left side (rows)'),
         ('B','Top (column headers)'),
         ('C','Either side — the result is the same'),
         ('D','Diagonally from top-left to bottom-right')],
        'B'
    ),
    # L6 — Embeddings — 3 questions
    (
        "In the CBOW model, the window parameter M = 2 means how many total context words are used as input?",
        [('A','2'), ('B','4'), ('C','1'), ('D','8')],
        'B'
    ),
    (
        "Why is the softmax function used in the Word2Vec output layer?",
        [('A','To reduce the embedding dimension'),
         ('B','To ensure output scores are positive and sum to 1 (a valid probability distribution)'),
         ('C','To speed up gradient descent by normalising gradients'),
         ('D','To prevent the centre word from appearing in the context')],
        'B'
    ),
    (
        "Skip-gram with Negative Sampling (SGNS) replaces the full softmax with what?",
        [('A','A sigmoid applied to k randomly sampled negative context words'),
         ('B','A Gaussian distribution over the entire vocabulary'),
         ('C','A lookup table that maps word IDs to fixed scores'),
         ('D','A cross-entropy loss over the top-1000 most frequent words')],
        'A'
    ),
    # L7 — PageRank — 2 questions
    (
        "In the PageRank stochastic adjacency matrix, columns represent SOURCE nodes. What must each column sum to?",
        [('A','0'), ('B','The PageRank score of that node'), ('C','1'), ('D','1/N')],
        'C'
    ),
    (
        "Without teleportation, what happens to PageRank scores when a 'spider trap' exists?",
        [('A','All scores converge to 1/N (uniform distribution)'),
         ('B','Scores for all nodes outside the trap converge to 0'),
         ('C','The power iteration diverges and never converges'),
         ('D','Only the dead-end nodes lose all their rank')],
        'B'
    ),
    # L8 — Graph — 2 questions
    (
        "Node v has degree 4 and there are 3 edges among its 4 neighbours. What is its clustering coefficient?",
        [('A','0.25'), ('B','0.50'), ('C','0.75'), ('D','1.00')],
        'B'  # 2*3/(4*3) = 6/12 = 0.5
    ),
    # L9 — NoSQL — 2 questions
    (
        "In the CAP theorem, a system configured as CP (Consistency + Partition Tolerance) must sacrifice:",
        [('A','Consistency during normal operation'),
         ('B','Availability when a network partition occurs'),
         ('C','Partition tolerance under high load'),
         ('D','Durability of committed writes')],
        'B'
    ),
    (
        "Which of the following BEST describes 'eventual consistency'?",
        [('A','All reads always return the most recent write immediately'),
         ('B','The system may return stale data briefly, but all nodes will eventually converge to the latest value'),
         ('C','Consistency is only guaranteed if fewer than 50% of nodes fail'),
         ('D','Data is consistent only within a single data centre, not across regions')],
        'B'
    ),
]

# ── LONG QUESTION DATA ───────────────────────────────────────────────────────

SCENARIO = """TechFlow is a smart city analytics company that collects data from 80,000 IoT devices \
across 15 cities. Every 200 milliseconds, sensors transmit readings including video feeds (MP4), \
JSON telemetry packets, CSV environmental logs, and free-text maintenance reports. \
The system ingests approximately 3.8 petabytes daily and must process emergency alerts \
within 2 seconds of detection. A recent audit revealed that 21% of sensor readings contain \
calibration drift errors or null values due to hardware faults. \
Parking sensors alone generate 40% of total data volume but contribute less than 2% of \
actionable city intelligence KPIs. During peak events (concerts, sports), data generation \
rates surge 8× and the types of relevant data sources shift entirely."""

MR_DOCS = [
    ("Doc1", ["spark", "data", "spark", "fast"]),
    ("Doc2", ["data", "lake", "spark", "data"]),
    ("Doc3", ["fast", "lake", "fast"]),
]

# BPE corpus: "aabb"×5, "aac"×3 — produces clean merge trace
BPE_CORPUS = [("aabb", 5), ("aac", 3)]
BPE_INITIAL_VOCAB = ['a', 'b', 'c', '</w>']
BPE_MERGES = [
    ("(a, a) → 'aa'", "Frequency: a-a = 5+3 = 8 (highest)", "aa bb </w> : 5\naa c </w> : 3"),
    ("(aa, b) → 'aab'", "Frequency: aa-b = 5 (highest)", "aab b </w> : 5\naa c </w> : 3"),
    ("(aab, b) → 'aabb'", "Frequency: aab-b = 5 (highest)", "aabb </w> : 5\naa c </w> : 3"),
]
BPE_FINAL_VOCAB = ['a', 'b', 'c', '</w>', 'aa', 'aab', 'aabb']  # size 7, was 4 (+3 merges)

# PageRank graph
PR_GRAPH = "Nodes: A, B, C.  Directed edges: A→C,  B→A,  B→C,  C→B"
PR_MATRIX_DATA = [
    ['', 'A (src)', 'B (src)', 'C (src)'],
    ['A (dst)', '0', '1/2', '0'],
    ['B (dst)', '0', '0', '1'],
    ['C (dst)', '1', '1/2', '0'],
]
PR_ITER1 = {
    'M_r': '[1/6, 1/3, 1/2]  (= [0.1667, 0.3333, 0.5000])',
    'r_new': 'β×(M·r) + (1−β)/3 = 0.8×[0.1667, 0.3333, 0.5000] + 0.0667\n'
             '        = [0.1333, 0.2667, 0.4000] + [0.0667, 0.0667, 0.0667]\n'
             '        = [0.2000, 0.3333, 0.4667]',
    'L1': '|0.2000−0.3333| + |0.3333−0.3333| + |0.4667−0.3333| = 0.1333 + 0 + 0.1333 = 0.2667',
}

# Graph analytics: undirected graph
# Nodes: 1,2,3,4,5   Edges: 1-2, 1-3, 2-3, 2-4, 3-5
GRAPH_ADJ = """
    1 — 2
    1 — 3
    2 — 3
    2 — 4
    3 — 5
"""
GRAPH_ANSWERS = {
    'degrees': {'1':2, '2':3, '3':3, '4':1, '5':1},
    'density': '5 / 10 = 0.50   [formula: 2|E| / N(N−1) = 10 / 20 = 0.50]',
    'cc_node2': '2×1 / (3×2) = 2/6 = 0.333   [m_2=1 edge (1-3); d_2=3]',
    'betweenness_2': '(1,4): 1/1=1  +  (3,4): 1/1=1  +  (4,5): 1/1=1  =  3.0   [normalised: 3/6 = 0.50]',
}

# Levenshtein: MART → KARMA
LEV_SOURCE = 'M A R T'
LEV_TARGET = 'K A R M A'
LEV_TABLE  = [
    ['',  '',  'M', 'A', 'R', 'T'],
    ['',   0,   1,   2,   3,   4 ],
    ['K',  1,   1,   2,   3,   4 ],
    ['A',  2,   2,   1,   2,   3 ],
    ['R',  3,   3,   2,   1,   2 ],
    ['M',  4,   3,   3,   2,   2 ],
    ['A',  5,   4,   3,   3,   3 ],
]

RDD_SKELETON = '''\
from pyspark import SparkContext
sc = SparkContext("local", "WordCount")

# TODO 1: Load "review.txt" as an RDD
lines = ___________________________

# TODO 2: Split each line into lowercase words (hint: use flatMap)
words = ___________________________

# TODO 3: Map each word to a (word, 1) pair
pairs = ___________________________

# TODO 4: Sum all values for the same word key
counts = ___________________________

# TODO 5: Return the top-5 most frequent words, sorted descending by count
top5 = ___________________________

print(top5)'''

RDD_ANSWERS = '''\
lines  = sc.textFile("review.txt")
words  = lines.flatMap(lambda line: line.lower().split())
pairs  = words.map(lambda w: (w, 1))
counts = pairs.reduceByKey(lambda a, b: a + b)
top5   = counts.takeOrdered(5, key=lambda x: -x[1])'''


# ══════════════════════════════════════════════════════════════════════════════
def build_exam(is_answer_key=False):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    suffix = ' — ANSWER KEY' if is_answer_key else ''

    # ── Cover ──
    t = doc.add_heading(f'DSAI4205 Big Data Analytics{suffix}', level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph('Mock Final Examination  |  PolyU · Dr. Ken Fong')
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(12)
    s.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Instructions box
    doc.add_paragraph()
    instr = doc.add_table(rows=1, cols=1)
    instr.style = 'Table Grid'
    cell = instr.rows[0].cells[0]
    cell_bg(cell, 'EFF6FF')
    lines_txt = [
        'Time allowed: 2 hours',
        'Section A: 22 multiple-choice questions (22 marks, 1 mark each)',
        'Section B: Long questions (78 marks)',
        'Programming accounts for approximately 10% of Section B marks.',
        'All RDD programming questions — DataFrame / SparkSQL will NOT appear in Section B.',
        'Calculators are NOT permitted.',
        'Show ALL working for computation questions.',
    ]
    cell.text = ''
    for i, line in enumerate(lines_txt):
        p = cell.add_paragraph()
        r = p.add_run(('INSTRUCTIONS\n' if i == 0 else '') + line)
        r.font.size = Pt(10)
        if i == 0:
            r.bold = True
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()
    doc.add_page_break()

    # ══ SECTION A ══
    doc.add_heading('SECTION A — Multiple Choice  (22 marks)', level=1)
    p = doc.add_paragraph('Circle the letter of the BEST answer for each question.')
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True
    if is_answer_key:
        p2 = doc.add_paragraph('Correct answers shown in bold green.')
        p2.runs[0].font.color.rgb = RGBColor(0x16, 0x65, 0x34)
        p2.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    for i, (stem, options, answer) in enumerate(MCQ_QUESTIONS, 1):
        mcq(doc, i, stem, options, answer, is_answer_key=is_answer_key)
        doc.add_paragraph()

    doc.add_page_break()

    # ══ SECTION B ══
    doc.add_heading('SECTION B — Long Questions  (78 marks)', level=1)
    p = doc.add_paragraph('Answer ALL questions. Show full working for all computation questions.')
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True
    doc.add_paragraph()

    # ── Q1: 6 V's (16 marks) ──────────────────────────────────────────────────
    q_heading(doc, 1, 'The 6 V\'s of Big Data  —  L1', 16)
    rule(doc)
    p = doc.add_paragraph()
    p.add_run('Read the following scenario and answer the questions below.\n').bold = True
    p.runs[0].font.size = Pt(10.5)
    box = doc.add_table(rows=1, cols=1)
    box.style = 'Table Grid'
    cell_bg(box.rows[0].cells[0], 'F0F9FF')
    box.rows[0].cells[0].text = SCENARIO
    for para in box.rows[0].cells[0].paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
    doc.add_paragraph()

    subq(doc, 'a', 'Identify SIX characteristics of Big Data (the 6 V\'s) evident in the scenario above. '
                   'For each V, quote or paraphrase the specific evidence from the passage.  (6 marks)')
    if is_answer_key:
        ans_tbl = doc.add_table(rows=7, cols=3)
        ans_tbl.style = 'Table Grid'
        headers = [('V', 'Evidence in passage', 'Proposed solution')]
        data = [
            ('Volume',      '3.8 petabytes daily; 80,000 devices',                         'Spark batch processing on HDFS / distributed storage'),
            ('Velocity',    'Every 200 ms; emergency alerts within 2 seconds',             'Spark Streaming / real-time pipeline'),
            ('Variety',     'MP4 video, JSON, CSV, free-text reports',                     'Data Lake; SparkSQL over mixed formats'),
            ('Veracity',    '21% of readings have calibration drift errors / null values', 'Data validation pipeline; anomaly detection'),
            ('Value',       'Parking sensors = 40% volume but <2% KPI contribution',       'Feature selection; deprioritise low-value streams'),
            ('Variability', 'Data rate surges 8× during events; relevant sources shift',   'Adaptive pipeline; periodic model retraining'),
        ]
        for ri, (v, ev, sol) in enumerate(headers + data):
            for ci, txt in enumerate([v, ev, sol]):
                c = ans_tbl.rows[ri].cells[ci]
                c.text = txt
                for para in c.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x16, 0x65, 0x34) if ri > 0 else RGBColor(0x1E, 0x3A, 0x5F)
                        if ri == 0:
                            run.bold = True
                if ri == 0:
                    cell_bg(c, '1E3A5F')
                    for para in c.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph()
    else:
        answer_box(doc, 10)

    subq(doc, 'b', 'For each V you identified, propose ONE specific technical solution. '
                   'Your solution must name the technology (e.g. Spark, HDFS, Kafka) and explain why it addresses that V.  (10 marks)')
    if not is_answer_key:
        answer_box(doc, 10)
    else:
        p = doc.add_paragraph()
        p.add_run('See solution column in table above.').font.color.rgb = RGBColor(0x16, 0x65, 0x34)
        p.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    doc.add_page_break()

    # ── Q2: MapReduce (18 marks) ───────────────────────────────────────────────
    q_heading(doc, 2, 'MapReduce  —  L2/L3', 18)
    rule(doc)
    p = doc.add_paragraph()
    p.add_run('Three documents are given below. A MapReduce word-count job is applied.\n').bold = True
    p.runs[0].font.size = Pt(10.5)
    for did, words in MR_DOCS:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        r = p2.add_run(f'{did}:  "{" ".join(words)}"')
        r.font.name = 'Courier New'
        r.font.size = Pt(10)
    doc.add_paragraph()

    subq(doc, 'a', 'State the key-value pair(s) emitted by the MAPPER for each document. '
                   'Show ALL emitted pairs.  (6 marks)')
    if is_answer_key:
        mr_ans = doc.add_table(rows=4, cols=2)
        mr_ans.style = 'Table Grid'
        cell_bg(mr_ans.rows[0].cells[0], '1E3A5F')
        cell_bg(mr_ans.rows[0].cells[1], '1E3A5F')
        mr_ans.rows[0].cells[0].text = 'Mapper'
        mr_ans.rows[0].cells[1].text = 'Emitted (key, value) pairs'
        for para in mr_ans.rows[0].cells[0].paragraphs + mr_ans.rows[0].cells[1].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        mr_data = [
            ('Doc1', '(spark,1) (data,1) (spark,1) (fast,1)'),
            ('Doc2', '(data,1) (lake,1) (spark,1) (data,1)'),
            ('Doc3', '(fast,1) (lake,1) (fast,1)'),
        ]
        for ri, (doc_id, pairs_str) in enumerate(mr_data, 1):
            mr_ans.rows[ri].cells[0].text = doc_id
            mr_ans.rows[ri].cells[1].text = pairs_str
            for ci in range(2):
                for para in mr_ans.rows[ri].cells[ci].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9.5)
                        run.font.name = 'Courier New'
                        run.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
        doc.add_paragraph()
    else:
        answer_box(doc, 6)

    subq(doc, 'b', 'Show the output AFTER the shuffle-and-sort phase. '
                   'List all (key, [values]) groups passed to the reducers.  (4 marks)')
    if is_answer_key:
        shuffle_ans = [
            '(data,  [1, 1, 1])',
            '(fast,  [1, 1, 1])',
            '(lake,  [1, 1])',
            '(spark, [1, 1, 1])',
        ]
        p_ans = doc.add_paragraph()
        p_ans.paragraph_format.left_indent = Cm(0.5)
        r_ans = p_ans.add_run('\n'.join(shuffle_ans))
        r_ans.font.name = 'Courier New'
        r_ans.font.size = Pt(10)
        r_ans.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
        doc.add_paragraph()
    else:
        answer_box(doc, 5)

    subq(doc, 'c', 'State the final output of the REDUCER. List all (word, count) pairs.  (4 marks)')
    if is_answer_key:
        ans_note(doc, 'data:3   fast:3   lake:2   spark:3')
        doc.add_paragraph()
    else:
        answer_box(doc, 3)

    subq(doc, 'd', 'Briefly explain why MapReduce writes intermediate results to HDFS after each stage '
                   'and why this makes it slower than Spark.  (4 marks)')
    if is_answer_key:
        ans_note(doc,
            'MapReduce writes to HDFS for fault tolerance: if a task fails, the system replays from the last checkpoint. '
            'HDFS replicates each write 3 times, so every mapper output is written 3× to disk — very high I/O cost. '
            'Spark avoids this by keeping RDDs in memory and using lineage for recovery, making it 10–100× faster '
            'for iterative workloads.')
        doc.add_paragraph()
    else:
        answer_box(doc, 4)

    doc.add_page_break()

    # ── Q3: PageRank (20 marks) ────────────────────────────────────────────────
    q_heading(doc, 3, 'PageRank  —  L7', 20)
    rule(doc)
    p = doc.add_paragraph()
    p.add_run('Consider the following directed graph:\n').bold = True
    p.runs[0].font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    r2 = p2.add_run(PR_GRAPH)
    r2.font.name = 'Courier New'
    r2.font.size = Pt(10.5)
    r2.bold = True
    doc.add_paragraph()

    subq(doc, 'a', 'Construct the column-stochastic adjacency matrix M. '
                   'State clearly which node corresponds to each row and column.  (5 marks)')
    if is_answer_key:
        dp_table(doc, PR_MATRIX_DATA)
        ans_note(doc, 'Out-degrees: A=1 (→C), B=2 (→A,→C), C=1 (→B). '
                      'Each column entry = 1/out-degree for nodes that are linked to, 0 otherwise. '
                      'All columns sum to 1. SOURCE at TOP (columns), DESTINATION at LEFT (rows).')
        doc.add_paragraph()
    else:
        answer_box(doc, 6)

    subq(doc, 'b', 'Using β = 0.8, perform ONE full power iteration starting from r⁰ = [1/3, 1/3, 1/3]. '
                   'Show the calculation for r¹ and the L1 norm change.  (9 marks)')
    if is_answer_key:
        pr_tbl = doc.add_table(rows=4, cols=2)
        pr_tbl.style = 'Table Grid'
        pr_data = [
            ('Step', 'Working'),
            ('M · r⁰', PR_ITER1['M_r']),
            ('r¹ = β(M·r⁰) + (1−β)/3', PR_ITER1['r_new']),
            ('L1 norm change', PR_ITER1['L1']),
        ]
        for ri, (k, v) in enumerate(pr_data):
            pr_tbl.rows[ri].cells[0].text = k
            pr_tbl.rows[ri].cells[1].text = v
            if ri == 0:
                cell_bg(pr_tbl.rows[ri].cells[0], '1E3A5F')
                cell_bg(pr_tbl.rows[ri].cells[1], '1E3A5F')
            for ci in range(2):
                for para in pr_tbl.rows[ri].cells[ci].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9.5)
                        run.font.name = 'Courier New'
                        run.font.color.rgb = (RGBColor(0xFF,0xFF,0xFF) if ri==0
                                              else RGBColor(0x16,0x65,0x34))
                        if ri == 0:
                            run.bold = True
        doc.add_paragraph()
    else:
        answer_box(doc, 7)

    subq(doc, 'c', 'What TWO problems does the teleportation parameter β solve? '
                   'Briefly explain why teleportation fixes each problem.  (4 marks)')
    if is_answer_key:
        ans_note(doc,
            '1) Spider traps: cycles absorb all rank. Teleportation lets the surfer escape with prob (1−β), '
            'redistributing rank to the rest of the graph.\n'
            '   2) Dead ends: nodes with no outlinks cause rank to "leak" to zero. Teleportation re-inserts '
            'leaked rank uniformly across all pages, preserving the total rank sum.')
        doc.add_paragraph()
    else:
        answer_box(doc, 4)

    subq(doc, 'd', 'Write the Google Matrix formula and identify each component.  (2 marks)')
    if is_answer_key:
        ans_note(doc, 'r = β·M·r + (1−β)/N · 1\n'
                      'β ≈ 0.85 = damping factor (prob of following a link)\n'
                      '(1−β)/N = uniform teleportation probability\n'
                      'M = column-stochastic adjacency matrix\n'
                      'N = total number of pages')
        doc.add_paragraph()
    else:
        answer_box(doc, 3)

    doc.add_page_break()

    # ── Q4: Graph Analytics (14 marks) ────────────────────────────────────────
    q_heading(doc, 4, 'Graph Analytics  —  L8', 14)
    rule(doc)
    p = doc.add_paragraph()
    p.add_run('Consider the following UNDIRECTED graph:\n').bold = True
    p.runs[0].font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    r2 = p2.add_run(GRAPH_ADJ.strip())
    r2.font.name = 'Courier New'
    r2.font.size = Pt(10.5)
    doc.add_paragraph()

    subq(doc, 'a', 'State the degree of each node (1 through 5).  (2 marks)')
    if is_answer_key:
        ans_note(doc, 'Node 1: 2   Node 2: 3   Node 3: 3   Node 4: 1   Node 5: 1')
        doc.add_paragraph()
    else:
        answer_box(doc, 2)

    subq(doc, 'b', 'Calculate the density of this graph. State the formula used.  (2 marks)')
    if is_answer_key:
        ans_note(doc, f'Density = 2|E| / N(N−1) = 2×5 / (5×4) = 10/20 = 0.50')
        doc.add_paragraph()
    else:
        answer_box(doc, 2)

    subq(doc, 'c', 'Compute the clustering coefficient of node 2. Show all working.  (4 marks)')
    if is_answer_key:
        ans_note(doc,
            'Neighbours of node 2: {1, 3, 4}\n'
            'Edges among neighbours: only 1–3 exists. So m₂ = 1.\n'
            'C(2) = 2·m₂ / (d₂·(d₂−1)) = 2×1 / (3×2) = 2/6 = 0.333')
        doc.add_paragraph()
    else:
        answer_box(doc, 5)

    subq(doc, 'd', 'Compute the betweenness centrality of node 2 (vertex betweenness only). '
                   'List all (source, target) pairs whose shortest path passes through node 2.  (4 marks)')
    if is_answer_key:
        ans_note(doc,
            'Pairs whose shortest path goes through node 2:\n'
            '  (1,4): 1→2→4  ✓  contribution 1/1 = 1\n'
            '  (3,4): 3→2→4  ✓  contribution 1/1 = 1\n'
            '  (4,5): 4→2→3→5  ✓  contribution 1/1 = 1\n'
            'All other pairs have a shortest path that avoids node 2.\n'
            'Betweenness(2) = 1 + 1 + 1 = 3   (normalised: 3 / [(5−1)(5−2)/2] = 3/6 = 0.50)')
        doc.add_paragraph()
    else:
        answer_box(doc, 6)

    subq(doc, 'e', 'Which type of betweenness centrality did Dr. Fong confirm is NOT tested in this exam?  (1 mark — 1 sentence)')
    if is_answer_key:
        ans_note(doc, 'Edge (arc) betweenness centrality is NOT tested. Only vertex betweenness centrality is in scope.')
        doc.add_paragraph()
    else:
        answer_box(doc, 1)

    subq(doc, 'f', 'Node 3 cannot reach node 4 via a direct edge. If you were computing closeness centrality '
                   'and found that two nodes in your graph were disconnected, what alternative measure would you use '
                   'and why?  (1 mark)')
    if is_answer_key:
        ans_note(doc,
            'Harmonic centrality = Σ 1/d(v,u) for all u≠v. When d(v,u) = ∞ (disconnected), the term contributes 0 '
            'instead of breaking the formula. Standard closeness breaks because ∞ makes the average infinite.')
        doc.add_paragraph()
    else:
        answer_box(doc, 2)

    doc.add_page_break()

    # ── Q5: NLP — BPE + Levenshtein (10 marks) ────────────────────────────────
    q_heading(doc, 5, 'NLP: BPE & Levenshtein  —  L5', 10)
    rule(doc)

    doc.add_paragraph().add_run('Part A — Byte Pair Encoding (5 marks)').bold = True

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    lines_bpe = [
        'The following corpus (with </w> end-of-word markers) has been character-split:',
        '',
        '    a  a  b  b  </w>  :  5 occurrences',
        '    a  a  c  </w>     :  3 occurrences',
        '',
        f'Initial vocabulary (size 4): {BPE_INITIAL_VOCAB}',
        'Target vocabulary size: 7',
    ]
    for line in lines_bpe:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        r = p2.add_run(line)
        r.font.size = Pt(10)
        if 'a  a' in line or 'Initial' in line or 'Target' in line:
            r.font.name = 'Courier New'
    doc.add_paragraph()

    subq(doc, 'a', 'Show the THREE merge steps required to reach the target vocabulary size of 7. '
                   'For each step: state the most frequent pair, the merged token, and the updated corpus.  (3 marks)')
    if is_answer_key:
        for i, (merge, freq, corpus_after) in enumerate(BPE_MERGES, 1):
            p_m = doc.add_paragraph()
            p_m.paragraph_format.left_indent = Cm(0.5)
            r1 = p_m.add_run(f'Merge {i}: {merge}')
            r1.bold = True
            r1.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
            r1.font.size = Pt(10)
            p_f = doc.add_paragraph()
            p_f.paragraph_format.left_indent = Cm(1.0)
            p_f.add_run(freq).font.size = Pt(9.5)
            p_c = doc.add_paragraph()
            p_c.paragraph_format.left_indent = Cm(1.0)
            r_c = p_c.add_run(corpus_after)
            r_c.font.name = 'Courier New'
            r_c.font.size = Pt(9.5)
            r_c.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
        ans_note(doc, f'Final vocabulary (size 7): {BPE_FINAL_VOCAB}')
        doc.add_paragraph()
    else:
        answer_box(doc, 8)

    subq(doc, 'b', 'If you try to tokenise the word "aad" using the final vocabulary, '
                   'what token would represent the character "d"? Explain why.  (2 marks)')
    if is_answer_key:
        ans_note(doc,
            '"d" maps to [UNK] (shown as "?") because "d" was never in the original corpus, '
            'so it never entered the base vocabulary. BPE can only represent characters/subwords '
            'it has seen during training. Unseen tokens become [UNK].')
        doc.add_paragraph()
    else:
        answer_box(doc, 3)

    # Levenshtein
    doc.add_paragraph().add_run('Part B — Levenshtein Distance (5 marks)').bold = True
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f'Compute the Levenshtein (edit) distance between  "{LEV_SOURCE}"  and  "{LEV_TARGET}".')
    r.font.size = Pt(10.5)
    doc.add_paragraph()

    subq(doc, 'c', 'Complete the dynamic programming table below. '
                   'Remember: SOURCE word at the TOP (columns), TARGET word at the LEFT (rows).  (4 marks)')
    if is_answer_key:
        dp_table(doc, LEV_TABLE)
        ans_note(doc, 'Edit distance = dp[5][4] = 3  '
                      '(Replace M→K, keep A, keep R, insert M before A at end — or other valid 3-operation sequence).')
        doc.add_paragraph()
    else:
        empty_table = [row[:] for row in LEV_TABLE]
        for ri in range(1, len(empty_table)):
            for ci in range(1, len(empty_table[ri])):
                empty_table[ri][ci] = ''
        dp_table(doc, empty_table)

    subq(doc, 'd', 'State the edit distance and give ONE possible sequence of operations.  (1 mark)')
    if is_answer_key:
        ans_note(doc, 'Distance = 3.  '
                      'E.g.: (1) substitute M→K, (2) keep A, (3) keep R, (4) insert M, (5) substitute T→A  '
                      '— or other valid path tracing the DP table.')
        doc.add_paragraph()
    else:
        answer_box(doc, 2)

    doc.add_page_break()

    # ── Q6: RDD Programming (10 marks) ────────────────────────────────────────
    q_heading(doc, 6, 'RDD Programming  —  L3  (~10% of exam)', 10)
    rule(doc)
    p = doc.add_paragraph()
    r = p.add_run(
        'Complete the PySpark RDD program below. '
        'A text file "review.txt" contains one sentence per line. '
        'Your program should count word frequencies and return the top-5 most frequent words.\n'
        'Fill in each TODO (2 marks each).')
    r.font.size = Pt(10.5)
    doc.add_paragraph()

    code_tbl = doc.add_table(rows=1, cols=1)
    code_tbl.style = 'Table Grid'
    cell_bg(code_tbl.rows[0].cells[0], '1E293B')
    code_para = code_tbl.rows[0].cells[0].paragraphs[0]
    code_run = code_para.add_run(RDD_SKELETON)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9.5)
    code_run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    doc.add_paragraph()

    if is_answer_key:
        ans_tbl = doc.add_table(rows=1, cols=1)
        ans_tbl.style = 'Table Grid'
        cell_bg(ans_tbl.rows[0].cells[0], '0F2E1A')
        ans_para = ans_tbl.rows[0].cells[0].paragraphs[0]
        ans_run = ans_para.add_run('ANSWERS:\n' + RDD_ANSWERS)
        ans_run.font.name = 'Courier New'
        ans_run.font.size = Pt(9.5)
        ans_run.font.color.rgb = RGBColor(0x86, 0xEF, 0xAC)
        doc.add_paragraph()
        notes = [
            'TODO 1: sc.textFile() — loads lines of a text file as an RDD of strings.',
            'TODO 2: flatMap(lambda line: line.lower().split()) — splits lines into words; flatMap flattens the result.',
            'TODO 3: map(lambda w: (w, 1)) — creates (word, 1) pairs for each word.',
            'TODO 4: reduceByKey(lambda a, b: a+b) — sums all 1\'s for each word key.',
            'TODO 5: takeOrdered(5, key=lambda x: -x[1]) — returns top-5 by count descending. '
                    'Alternative: sortBy(lambda x: -x[1]).take(5)',
        ]
        for note in notes:
            p_n = doc.add_paragraph()
            p_n.paragraph_format.left_indent = Cm(0.5)
            r_n = p_n.add_run(note)
            r_n.font.size = Pt(9.5)
            r_n.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
    else:
        answer_box(doc, 1)

    return doc


# ══════════════════════════════════════════════════════════════════════════════
exam_doc = build_exam(is_answer_key=False)
exam_doc.save(str(OUT))
print(f'Saved exam: {OUT}')

key_path = OUT.parent / 'DSAI4205_Mock_Exam_ANSWERS.docx'
key_doc = build_exam(is_answer_key=True)
key_doc.save(str(key_path))
print(f'Saved key:  {key_path}')

# Quick stats
from docx import Document as DDoc
for path in [OUT, key_path]:
    d = DDoc(str(path))
    non_empty = [p for p in d.paragraphs if p.text.strip()]
    print(f'{path.name}: {len(non_empty)} paragraphs, {len(d.tables)} tables')
