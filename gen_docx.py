# gen_docx.py — regenerate DSAI4205_Exam_Review.docx with full content + ELI5
import re
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HTML = Path('C:/Users/User/Downloads/4205/DSAI4205_Exam_Review.html')
DOCX = Path('C:/Users/User/Downloads/4205/DSAI4205_Exam_Review.docx')

# ── ELI5 lookup ──────────────────────────────────────────────────────────────
ELI5 = {
    # L1
    "The 6 V's of Big Data": "Imagine you have a HUGE box of Lego bricks that keeps growing every second, comes in all shapes and colours, and sometimes has broken pieces mixed in. Big Data is like that — Volume (how many bricks), Velocity (how fast they arrive), Variety (all the different shapes), Veracity (some are fake/broken), Value (only some bricks are useful), Variability (the shapes keep changing).",
    "Three Sources of Big Data": "Data comes from three places: things people say online (social), machines that automatically log stuff (machine), and records of buying/selling (transactional). Think Twitter posts, your car's GPS sensor, and your receipt at 7-Eleven.",
    "Data Analytics Process (5 Steps)": "It's like solving a mystery: first decide WHAT you're trying to find out, then collect clues (data), then look at the clues carefully, then make your best guess, then tell everyone what you found.",
    "Why Traditional Tools Fail": "Your laptop can handle a class assignment, but not all of Google's searches. Traditional databases are like a single filing cabinet — great for small offices, terrible for a skyscraper full of files.",
    "Types of Analytics": "Descriptive = reading your school report. Diagnostic = asking your teacher WHY you failed. Predictive = guessing what grade you'll get next term. Prescriptive = your teacher telling you exactly what to study to pass.",
    "AI History — Key Dates": "Key milestones: ELIZA (1964–67) was the first chatbot — a clever script that made you think you were talking to a therapist. IBM Deep Blue beat chess world champion Kasparov in 1997. AlphaGo beat Go champion Lee Sedol in 2016 — Go has more possible moves than atoms in the universe.",
    "Big Data Tools Landscape (L1)": "Think of it as a toolbox: Hadoop/HDFS stores huge files across many computers, Spark processes them fast, Hive lets you query with SQL-like language, and Kafka handles real-time data streams.",
    "Exam Practice — Netflix mapped to 4 V's": "Netflix: Volume = millions of users watching simultaneously; Velocity = real-time viewing data; Variety = movies, shows, metadata, ratings; Veracity = rating biases and missing data.",

    # L2
    "Scale Up vs Scale Out": "Scale Up = buy a bigger, more powerful computer (expensive, has limits). Scale Out = buy many cheap computers and share the work (like hiring more people instead of one super-employee).",
    "Hashing Strategy for Sharding": "When you have millions of rows, you split them across machines using a hash function — like assigning students to classrooms by the first letter of their surname. Same letter always goes to the same room.",
    "Dask Collections": "Dask gives you Pandas DataFrames and NumPy arrays that work on datasets too big to fit in RAM, by splitting them into chunks and processing one chunk at a time.",
    "MapReduce": "Map = give every worker one small task (count words on one page). Reduce = collect all their answers and combine (total word count across all pages). Like a class counting words in a book — each student does one chapter, then you add everyone's totals.",
    "HDFS (Hadoop Distributed File System)": "HDFS is like Google Drive but for a cluster. It splits files into 128MB blocks and stores 3 copies of each block on different machines. If one machine dies, the data is still safe on the other two.",
    "Limitations of MapReduce / Why Spark?": "MapReduce writes intermediate results to disk after every step — like doing homework, printing it, shredding it, then retyping it for the next step. Spark keeps results in RAM instead, making it 10-100x faster.",
    "Dask — Block Algorithms": "Dask processes one chunk at a time and keeps only the necessary partial result in memory. Like washing a mountain of dishes by doing one sink-full at a time instead of needing a bathtub big enough for all of them.",
    "Dask Task Graph": "Before doing any work, Dask draws a flowchart of ALL the steps needed. Then it executes the whole plan at once efficiently, rather than doing each step one-by-one naively.",
    "Dask Chunk Specification (4 ways)": "You can tell Dask how to split your data: by a fixed number (e.g., 1000 rows), by target size in MB, by a specific tuple of chunk sizes, or let Dask figure it out automatically.",
    "MapReduce — Detailed Walkthrough": "Step 1: Split data. Step 2: Map (each worker processes its split). Step 3: Shuffle (group same keys together). Step 4: Reduce (combine grouped values). Like a kitchen brigade — each chef has a role, food moves through the line.",
    "Dask Distributed Client": "The Dask Client is like a project manager who assigns tasks to workers, monitors progress, and collects results. You call `client.compute()` to actually run the task graph.",
    "Zarr — Chunked Array Storage": "Zarr stores large arrays (like satellite images) chunked into pieces, compressed, and saveable to disk or cloud. Think of it as a highly efficient filing system for multidimensional data.",
    "Dask — Eager vs Lazy Execution": "Lazy = Dask builds a plan but doesn't run it yet (like writing a grocery list). Eager = it runs immediately. Call `.compute()` to trigger lazy execution.",
    "Dask — Important API Details": "Key facts: `ddf.compute()` triggers execution; `ddf.persist()` keeps result in distributed memory; `ddf.visualize()` draws the task graph; `dask.delayed()` wraps any function to be lazy.",
    "Dask Performance Benchmark": "Dask vs single-threaded: for large datasets Dask is dramatically faster, but for small datasets the overhead of task scheduling can make it SLOWER. Always benchmark.",
    "Chunk Size Trade-offs": "Chunks too small = too much overhead scheduling millions of tiny tasks. Chunks too big = out-of-memory errors. Sweet spot: 10MB–1GB per chunk depending on task.",
    "Dask Advanced API": "Dask also supports Bag (for unstructured data like JSON), Delayed (wrap any Python function), and Futures (for real-time async computation with the distributed scheduler).",

    # L3
    "Core Concepts": "Spark keeps your data in memory (RAM) instead of writing to disk between steps. It creates a DAG (flowchart) of all operations and only runs them when you ask for results.",
    "RDD — Data Lineage & Fault Tolerance": "Every RDD remembers how it was created (lineage). If a partition is lost, Spark recomputes ONLY that partition by replaying the lineage — like redoing just the lost pages of homework, not the whole book.",
    "RDD — Granularity of Data Flow": "RDDs process data at the element level (one item at a time), while DataFrames process in columnar batches, which is faster for structured data.",
    "Transformations vs Actions": "Transformations (map, filter) are lazy — they just add steps to the plan. Actions (collect, count, save) actually trigger execution. Think of transformations as writing a recipe and actions as actually cooking.",
    "Narrow vs Wide Transformations": "Narrow = each output partition depends on ONE input partition (fast, no data movement). Wide = output depends on MULTIPLE input partitions — requires a shuffle (data moves between machines, expensive).",
    "Key DataFrame Operations": "select, filter, groupBy, agg, join, withColumn — these are your bread and butter. DataFrames are like supercharged Pandas that run on a cluster.",
    "Extended RDD Operations": "flatMap flattens lists-of-lists into one list. reduceByKey sums/combines values per key. sortByKey sorts by key. These are the classic building blocks.",
    "Key-Value RDD Joins": "join, leftOuterJoin, rightOuterJoin, fullOuterJoin — same semantics as SQL joins but on RDDs of (key, value) tuples.",
    "map() vs mapValues() vs flatMapValues()": "map() transforms the whole tuple. mapValues() only transforms the value (keeps the key). flatMapValues() transforms value into multiple values and flattens. Use mapValues() when you want to preserve partitioning.",
    "More DataFrame Patterns": "Window functions (rank, lag, lead over partitions), pivot tables, explode for arrays, struct/map types — advanced patterns for real analytics.",
    "Pandas — More Patterns": "groupby + agg, merge, apply, pivot_table — core Pandas patterns tested in tutorials.",
    "SparkSession Setup": "`SparkSession.builder.appName('x').getOrCreate()` is always the first line. SparkSession is your entry point to all Spark functionality.",
    "Catalyst Optimizer (DataFrame query engine)": "Catalyst is Spark's query planner. It rewrites your code into the most efficient physical plan — like a GPS recalculating the fastest route. It does predicate pushdown, column pruning, and join reordering.",
    "Nested Data Structures in DataFrames": "Spark DataFrames can hold arrays, maps, and structs inside columns. Use `explode()` to unnest arrays, `getItem()` to access map values.",
    "Partitioning & Join Optimisation": "Broadcast join = copy a small table to every worker (avoids shuffle). Repartition by key before joining large tables to co-locate matching keys on the same machine.",
    "Inverted Index (Take-Home Exercise)": "An inverted index maps each word to the list of documents containing it — like a book's index. Build with flatMap (word → (word, docID)) then groupByKey.",
    "Spark — Five Components": "SparkCore (RDD engine), SparkSQL (DataFrame/SQL), SparkStreaming (real-time), MLlib (machine learning), GraphX (graph computation).",
    "Spark Architecture Terminology": "Driver = coordinator process. Executors = worker processes on cluster nodes. Tasks = units of work. Stages = groups of tasks separated by shuffles.",
    "Fault Tolerance — Driver vs Worker Failure": "Worker failure → Spark re-schedules that worker's tasks on other workers (using lineage). Driver failure → the whole application must restart (driver is single point of failure).",
    "Cache Management": "`.cache()` stores RDD/DataFrame in memory. `.persist(StorageLevel.DISK_ONLY)` writes to disk. `.unpersist()` frees memory. Cache iteratively-used RDDs to avoid recomputation.",
    "Spark Log Levels": "Log levels: ALL < DEBUG < INFO < WARN < ERROR < FATAL < OFF. Set with `sc.setLogLevel('WARN')` to reduce noise.",

    # L4
    "SparkSQL": "SparkSQL lets you write normal SQL queries against DataFrames. You register a DataFrame as a temp view, then query it with `spark.sql('SELECT ...')`.",
    "Data Formats": "Parquet = columnar format, great for analytics (reads only needed columns). ORC = similar to Parquet. Avro = row-based, good for streaming. JSON/CSV = human-readable but slow. Always prefer Parquet for Spark.",
    "Hive vs SparkSQL": "Hive uses MapReduce underneath (slow). SparkSQL uses Spark (fast). Both understand HiveQL syntax. SparkSQL is 10-100x faster for interactive queries.",
    "Data Lake vs Data Warehouse": "Data Lake = dump everything raw (schema-on-read, any format). Data Warehouse = structured, cleaned, schema enforced (schema-on-write). Lake is flexible, Warehouse is fast for business queries.",
    "HiveQL — Key Commands": "CREATE TABLE, LOAD DATA, INSERT INTO, SELECT with GROUP BY, JOIN — nearly identical to SQL. Key difference: LOAD DATA copies files into Hive's warehouse directory.",
    "Hive Buckets (Bucketing)": "Bucketing splits data within partitions into fixed-size buckets using a hash of a column. Enables efficient same-key joins by ensuring matching keys land in the same bucket.",
    "Hive — Origin & History": "Hive was created at Facebook to let non-engineers query Hadoop with SQL. It translates HiveQL into MapReduce jobs. Now mostly replaced by SparkSQL for performance.",
    "Hive 3-Level Data Model": "Database → Table → Partition. Partitions are physical subdirectories (e.g., /date=2024/). Querying with WHERE on partition column skips reading other partitions (partition pruning).",
    "Hive Architecture — 4 Components": "Metastore (schema registry), Driver (query planner), Compiler (translates to execution plan), Execution Engine (runs the plan on Hadoop or Spark).",
    "HiveQL — Window Functions": "RANK(), ROW_NUMBER(), LAG(), LEAD() OVER (PARTITION BY ... ORDER BY ...) — compute running totals, rankings, and comparisons to previous/next rows without collapsing groups.",
    "HiveQL — LOAD DATA OVERWRITE": "LOAD DATA INPATH overwrites existing table data. LOAD DATA LOCAL INPATH copies from local filesystem. Without LOCAL, it moves from HDFS.",
    "Catalyst Optimizer — Advanced Rewrites": "Catalyst applies: predicate pushdown (filter early), column pruning (read only needed columns), join reordering, and constant folding. These happen automatically.",

    # L5
    "Linguistics Framework (NLP Subfields)": "NLP breaks language into layers: Phonology (sounds), Morphology (word parts like -ing, -ed), Syntax (grammar rules), Semantics (meaning), Pragmatics (context/intent). Most NLP tasks operate at morphology and syntax level.",
    "NLP Applications": "Spam filtering, sentiment analysis, machine translation, chatbots, search engines, speech recognition — everywhere text or speech is processed automatically.",
    "NLTK Library Overview": "NLTK is Python's classic NLP toolbox. It has tokenizers, stemmers, POS taggers, parsers, WordNet access, and corpora. Think of it as a Swiss Army knife for text.",
    "Text Normalisation — Regex Patterns": "Normalisation = making text consistent. Lowercase everything, remove punctuation with regex `[^a-z0-9\\s]`, collapse multiple spaces. This ensures 'The' and 'the' are the same token.",
    "Chunking — Noun Phrase Extraction": "Chunking uses POS tags and grammar rules to pull out noun phrases. Grammar like `NP: {<DT>?<JJ>*<NN>}` means: optional determiner + zero or more adjectives + noun = a noun phrase.",
    "POS-Tag Filtering (Keep Content Words Only)": "After POS tagging, keep only Nouns (NN*), Verbs (VB*), Adjectives (JJ*), Adverbs (RB*) — these carry meaning. Drop articles, prepositions, pronouns (stop words).",
    "Text Preprocessing Pipeline": "Raw text → lowercase → remove punctuation → tokenise → remove stopwords → stem/lemmatise → feature vector. Each step strips away noise.",
    "Stemming vs Lemmatization": "Stemming = chop the end off (running → run, studies → studi — sometimes wrong). Lemmatization = look up the real base form (studies → study). Lemmatization is smarter but slower.",
    "Text Feature Representations": "Bag of Words = count how many times each word appears (ignores order). TF-IDF = weight words by how unique they are across documents. Word2Vec = represent words as vectors capturing meaning.",
    "TF-IDF — Full Formula with Worked Example": "TF = (word count in doc) / (total words in doc). IDF = log(total docs / docs containing word). TF-IDF = TF × IDF. High TF-IDF = word is common in THIS doc but rare across all docs = important signal.",
    "Spark ML NLP Pipeline": "Pipeline([Tokenizer, StopWordsRemover, CountVectorizer/HashingTF, IDF, StringIndexer, LogisticRegression]) — chain transformers and estimators, fit once, transform the whole dataset.",
    "UDFs (User-Defined Functions)": "UDFs let you apply custom Python functions to DataFrame columns. `@udf(returnType=StringType())` decorator. Warning: UDFs are slower than built-in functions — avoid when possible.",
    "countByKey() vs countByValue()": "countByKey() works on (key, value) pairs — counts per key. countByValue() works on plain RDDs — counts each distinct element. Both return dicts but are actions (eager).",
    "Levenshtein (Edit) Distance": "Minimum number of single-character edits (insert, delete, substitute) to turn one word into another. 'kitten' → 'sitting' = 3 edits. Used in spell-checking and fuzzy matching.",
    "Byte Pair Encoding (BPE) — Subword Tokenization": "BPE starts with characters, then repeatedly merges the most frequent character pair into a new token. 'low lower lowest' → learns 'low', 'er', 'est' as tokens. Handles unknown words by splitting into known subwords.",
    "Corpus & Annotation": "A corpus is a collection of text. Annotation = humans labelling it (POS tags, named entities, sentiment). Supervised NLP models need annotated corpora for training.",
    "Penn Treebank — 14 Phrasal / Clause Categories": "The Penn Treebank uses tags like NP (noun phrase), VP (verb phrase), PP (prepositional phrase), S (sentence), SBAR (subordinate clause). These are used for syntactic parsing.",
    "NER — Named Entity Categories": "Named Entity Recognition identifies: PERSON, ORGANIZATION, LOCATION, DATE, TIME, MONEY, PERCENT, FACILITY, GPE (geopolitical entity). NLTK's `ne_chunk()` does this.",
    "TF-IDF — Two TF Normalisation Formulas": "Raw TF = count / total words. Log-normalised TF = 1 + log(count). Log normalisation reduces the impact of very frequent terms.",
    "Sparse Vector Format": "TF-IDF vectors are sparse (mostly zeros). Stored as (index, value) pairs or scipy `csr_matrix` instead of full arrays. Saves huge amounts of memory.",

    # L6
    "Distributional Hypothesis": "Words that appear in similar contexts have similar meanings. 'Dog' and 'puppy' often appear near 'bark', 'walk', 'leash' — so they get similar vector representations. This is the foundation of all word embeddings.",
    "Semantic Relationships in Embeddings": "In embedding space, similar words are close together. 'King' and 'Queen' are close. 'Paris' and 'France' have the same relationship as 'Berlin' and 'Germany'. Geometry encodes meaning.",
    "Word Vector Analogy Arithmetic": "king - man + woman ≈ queen. This works because word2vec encodes semantic relationships as vector directions. Royalty is a direction in the vector space.",
    "Special Tokens": "[PAD] = padding to make sequences same length. [UNK] = unknown word not in vocabulary. [CLS] = start-of-sequence in BERT. [SEP] = separator between sentences.",
    "Word Embeddings": "Instead of one-hot vectors (huge and sparse), embeddings are dense 50-300 dimensional vectors. Each dimension captures some latent semantic feature. Learned from large text corpora.",
    "Word2Vec — Two Architectures": "Skip-gram: given a word, predict its neighbours. CBOW: given neighbours, predict the centre word. Skip-gram works better for rare words; CBOW is faster to train.",
    "Skip-gram: Objective Function": "Maximise the probability of observing each context word given the centre word, summed over all (centre, context) pairs in the corpus. Uses softmax over the whole vocabulary.",
    "Skip-gram with Negative Sampling (SGNS)": "Full softmax over vocabulary is too slow. Negative sampling: for each real context word, sample k random 'negative' words. Train a binary classifier: real context vs noise. Much faster.",
    "CBOW — Loss Function & Detail": "CBOW averages the context word vectors, passes through a hidden layer, and predicts the centre word. Loss = cross-entropy between predicted and actual centre word.",
    "Skip-gram — Three Formal Assumptions": "1) Context is a fixed window of k words either side. 2) Each context word is independent (bag of context, ignores order). 3) Word2Vec embeddings are static — no context sensitivity (unlike BERT).",
    "GloVe — Global Vectors (Stanford)": "GloVe = Global Vectors. Builds a word co-occurrence matrix from the ENTIRE corpus, then factorises it. Captures global statistics better than Word2Vec's local window approach.",
    "CBOW Neural Architecture": "Input: average of context word one-hot vectors. Hidden: weight matrix W (V × N). Output: softmax over vocabulary. Trained by backprop. N = embedding dimension.",
    "Text Similarity Algorithms": "Cosine similarity: angle between vectors (ignores magnitude). Jaccard: intersection/union of word sets. Edit distance: character-level edits. Sorensen-Dice: 2|A∩B|/(|A|+|B|).",
    "Neural Network Components (Review)": "Input layer → hidden layers with activation functions (ReLU, sigmoid) → output layer. Weights updated by gradient descent via backpropagation. Loss function measures error.",
    "CBOW — Numerical Example": "With 2-word context window and 3-word vocab, average the two context word embeddings, multiply by output weight matrix W', apply softmax. Highest probability = predicted word.",
    "From Word2Vec to Transformers (conceptual)": "Word2Vec gives STATIC embeddings (one vector per word regardless of context). BERT/Transformers give CONTEXTUAL embeddings — 'bank' gets different vectors in 'river bank' vs 'bank account'.",
    "Text Similarity - Full Numerical Examples": "Cosine sim between (1,0,1) and (0,1,1) = (0+0+1)/(√2 × √2) = 0.5. Jaccard between {a,b,c} and {b,c,d} = 2/4 = 0.5. Work through with actual numbers.",
    "Cosine Similarity - Worked Numerical Example": "cos(θ) = (A·B) / (|A| × |B|). Dot product = sum of element-wise products. Magnitude = square root of sum of squares. Result in [-1, 1]. 1 = identical direction, 0 = orthogonal, -1 = opposite.",
    "Word Analogy Arithmetic - Coordinate Trace": "v(king) - v(man) + v(woman) = result vector. Find the word in vocabulary whose embedding is nearest to result. This is vector subtraction then nearest-neighbour search.",
    "Word2Vec - Practical Embedding Facts": "Typical dimension: 100–300. Window size: 2–10 words. Training on billions of words. Pre-trained models: Google News Word2Vec (3M words, 300d), GloVe-6B (Wikipedia+Gigaword).",
    "Skip-gram Weight Matrix Dimensions": "Input weight matrix W: V × N. Output weight matrix W': N × V. V = vocabulary size, N = embedding dimension. The rows of W are the final word embeddings.",
    "GloVe - Weighting Function and Hyperparameters": "Weighting function f(x) = (x/x_max)^α if x < x_max else 1. x_max = 100, α = 0.75 by default. Prevents very frequent co-occurrences from dominating the loss.",
    "Levenshtein DP - Cell-by-Cell Fill Logic": "DP table: rows = characters of word1 (+empty), cols = characters of word2 (+empty). Cell[i][j] = min(delete, insert, substitute). If characters match, substitute cost = 0.",

    # L7
    "Web as a Graph": "The web is a directed graph: web pages are nodes, hyperlinks are directed edges. PageRank computes the importance of each page based on the structure of this graph.",
    "Core Idea": "A page is important if MANY pages link to it, AND those linking pages are themselves important. It's recursive — importance flows through links like water flowing downhill.",
    "Random Walk Interpretation": "Imagine a random surfer clicking links forever. PageRank = probability that the surfer is on a given page at any moment. Pages with more/better inlinks = visited more often.",
    "Stochastic Adjacency Matrix M": "M is the column-normalised adjacency matrix. Each column sums to 1 (probability distribution over outlinks). PageRank is the eigenvector of M with eigenvalue 1.",
    "Power Iteration": "Start with equal probability for all pages (1/N each). Multiply by M repeatedly. After ~50 iterations, the vector converges to the stationary distribution = PageRank scores.",
    "Problems and Fixes": "Spider traps (cycles of pages with no outlinks to rest of web) absorb all rank. Dead ends (pages with no outlinks) leak rank. Fix: add teleportation (random jump with probability 1-β).",
    "PageRank at Scale": "For billions of pages, store M as sparse matrix. Process in blocks. Use MapReduce: each node sends its rank divided equally to all its outlink targets.",
    "Google Matrix (The Full Formula)": "r = β × M × r + (1-β)/N × 1. β ≈ 0.85 (damping factor). (1-β)/N is the teleportation term — with prob 1-β, surfer jumps to any random page. Solves both spider trap and dead-end problems.",
    "Efficient Sparse Implementation (Tutorial 7 formula)": "For dead ends: add leaked rank (sum of dead-end ranks / N) to all nodes. Then apply: r_new = β × M × r + ((1-β) + leaked) / N × 1.",
    "Spider Trap — Numerical Example": "Two nodes A→B, B→A, B→B (B is a trap). Without teleportation, all rank flows into B. With β=0.8: rank distributes as B gets ~83%, A gets ~17% due to teleportation.",
    "Dead End — Numerical Example": "Node C with no outlinks. Without fix, C absorbs rank but gives nothing back — rank 'leaks'. Fix: treat dead ends as linking to ALL nodes equally (redistributes leaked rank).",
    "PageRank Convergence — Numerical Trace (β=0.8)": "Track r vector at each iteration. Compute L1 norm change: sum(|r_new - r_old|). Stop when change < threshold (e.g., 0.001). Typically converges in 50-100 iterations.",
    "Worked Example (3-node graph)": "With 3 nodes, write out M explicitly, multiply β×M×r+(1-β)/3, iterate until convergence. Know how to do this by hand for small graphs.",
    "WWW History": "Tim Berners-Lee invented the World Wide Web in 1989 at CERN. First website went live in 1991. The web grew exponentially — PageRank (1998) was the insight that made Google's search far better than competitors.",
    "Links as Votes - Trust Propagation Intuition": "A link from a high-quality site (like Wikipedia) is worth much more than a link from a random blog. PageRank propagates trust recursively through the link graph.",
    "PageRank Score Motivating Example": "If a page has 5 links pointing to it from pages with high PageRank, it gets more rank than a page with 100 links from pages with very low rank. Quality beats quantity.",
    "Flow Equations - Gaussian Elimination (small graphs only)": "For tiny graphs (3-4 nodes), write flow equations: r(A) = sum over B linking to A of r(B)/out_degree(B). Solve as a linear system. Only feasible for toy examples.",
    "Simple Cycle - Stationary Distribution": "For a simple directed cycle A→B→C→A, the stationary distribution is uniform (1/3 each) — rank distributes equally because all nodes are equivalent.",
    "Dead-End Full Algorithm - Leaked Rank Re-insertion": "Algorithm: 1) compute M×r for non-dead-end nodes; 2) sum rank of dead-end nodes; 3) add leaked_rank/N to all nodes; 4) apply teleportation; 5) normalise.",
    "Memory at Scale + Sparse Matrix Trick": "Store only non-zero entries of M as (from, to, 1/out_degree) triples. For the web, M is 99.9%+ zeros. Sparse representation reduces memory from O(N²) to O(edges).",
    "MapReduce PageRank - Pseudocode": "Map: for each (page, [links, rank]): emit (page, rank/|links|) to each linked page, emit (page, [links, 0]) to preserve structure. Reduce: sum all received contributions + apply damping.",
    "Convergence Trace - L1 Norm Per Iteration": "After each power iteration, compute Σ|r_new[i] - r_old[i]|. This L1 norm should decrease monotonically. Stop when it drops below threshold (e.g., 0.001).",

    # L8
    "Graph Types": "Undirected (Facebook friends — symmetric). Directed (Twitter follows — A can follow B without B following A). Weighted (road distances). Bipartite (users × items). Multigraph (multiple edges between same nodes).",
    "Graph Basics": "G = (V, E). V = vertices/nodes. E = edges. Degree = number of edges on a node. Path = sequence of nodes. Connected = there's a path between every pair. Diameter = longest shortest path.",
    "Clustering Coefficient": "Local CC = (actual triangles through node) / (possible triangles through node). Measures how 'cliquey' a node's neighbours are. High CC = tight-knit community. Global CC = average over all nodes.",
    "Centrality Measures": "Degree = how many connections. Betweenness = how often on shortest paths (a bridge). Closeness = how quickly can reach everyone. Eigenvector = connected to well-connected nodes (like PageRank for general graphs).",
    "Real-World Case Study: MSN Messenger Network (2006)": "180M nodes, 1.3B edges. Average path length = 6.6 (six degrees of separation). Clustering coefficient = 0.11. Follows scale-free (power law) degree distribution — most nodes have few connections, few have many.",
    "Centrality Limitations — Krackhardt's Kite Graph": "Designed to show all centrality measures pick different 'most important' nodes. Degree centrality picks the popular hub. Betweenness picks the broker who bridges groups. Closeness picks who can spread info fastest.",
    "Network Properties (Small-World & Scale-Free)": "Small-World: high clustering + short paths (like real social networks). Scale-Free: power-law degree distribution (most nodes have few connections, a few hubs have many — like web pages, airports).",
    "Graph Connectivity": "Connected (undirected): path exists between all pairs. Strongly connected (directed): path exists in BOTH directions between all pairs. Weakly connected (directed): ignoring direction, it's connected.",
    "Betweenness Centrality — Worked Example": "For each pair (s,t): count all shortest paths σ(s,t). Count how many pass through node v: σ(s,t|v). Betweenness(v) = Σ σ(s,t|v)/σ(s,t). High betweenness = critical bridge/broker.",
    "Degree Distribution": "P(k) = probability that a random node has degree k. Random graphs: Poisson distribution. Scale-free networks: power law P(k) ∝ k^(-γ) with γ typically 2-3. Power law = very unequal, hub-and-spoke structure.",
    "Directed vs Undirected - Max Edges and Density": "Undirected max edges: N(N-1)/2. Directed max edges: N(N-1). Density = actual edges / max edges. Sparse graph: density << 1. Dense graph: density close to 1.",
    "Degree Distribution - Formal Definition": "P(k) = |{v : degree(v) = k}| / |V|. The degree sequence is the sorted list of all node degrees. For directed graphs, separate in-degree and out-degree distributions.",
    "Path Definition - Can Revisit": "A walk can revisit nodes/edges. A path cannot revisit nodes. A trail cannot revisit edges. Shortest path = path with minimum number of edges (or minimum weight for weighted graphs).",
    "MSN Messenger - Full Statistics": "180M nodes, 1.3B edges, average degree 14.4, diameter 29 (but average path length 6.6), clustering coefficient 0.11, power-law degree distribution. Classic small-world + scale-free network.",
    "In-Degree Centrality = Majority Voting": "In-degree counts who points TO you. High in-degree = many people reference/follow you. This is why Twitter follower count = a form of in-degree centrality. Correlates with influence.",
    "Closeness Centrality - Disconnected Graph Failure": "Closeness = 1 / (average shortest path to all others). Problem: in disconnected graphs, some pairs have infinite distance → closeness = 0 (meaningless). Fix: use harmonic centrality = Σ 1/d(u,v).",
    "Eigenvector Centrality - 4-Step Power Iteration Pseudocode": "1) Init all scores = 1. 2) New score of v = sum of scores of v's neighbours. 3) Normalise. 4) Repeat until convergence. High eigenvector centrality = connected to other high-scoring nodes (recursive).",
    "Which Centrality to Use?": "Degree: who is popular/active. Betweenness: who is a broker/bridge. Closeness: who can spread info fastest. Eigenvector/PageRank: who is connected to important nodes. Choose based on the question.",
    "Edge Betweenness - Forward BFS + Backward Fractional Flow Algorithm": "1) BFS forward from source s to find all shortest paths and their counts. 2) Backward pass: assign fractional flow (1/paths) to each edge on shortest paths. 3) Sum over all (s,t) pairs. High edge betweenness = critical bridge edge.",

    # L9
    "ACID Properties (Traditional RDBMS)": "Atomicity (all-or-nothing transaction), Consistency (database rules always satisfied), Isolation (concurrent transactions don't interfere), Durability (committed data survives crashes). Traditional databases guarantee all four.",
    "CAP Theorem": "In a distributed system, you can only guarantee 2 of 3: Consistency (all nodes see the same data), Availability (every request gets a response), Partition tolerance (system works despite network splits). Networks always partition, so choose CA or CP.",
    "Strong vs Eventual Consistency": "Strong: every read returns the most recent write. Eventual: if you stop writing, eventually all copies will be consistent (but you might read stale data in the meantime). Strong is safe but slow; eventual is fast but risky.",
    "BASE Model (NoSQL alternative to ACID)": "Basically Available (system is always responsive), Soft state (data may change over time without input — due to replication), Eventually consistent. The NoSQL alternative to ACID for distributed systems.",
    "NoSQL Data Models": "Key-Value (Redis, DynamoDB): simple get/put by key. Document (MongoDB): JSON documents with flexible schema. Wide Column (Cassandra, HBase): rows with dynamic columns. Graph (Neo4j): nodes and edges for relationship data.",
    "Transactions & OLTP": "OLTP = Online Transaction Processing — many small, fast transactions (bank transfers, orders). Needs ACID. Contrast with OLAP (analytics queries on historical data — batch, read-heavy).",
    "Key-Value Store Operations": "GET(key) → value. PUT(key, value). DELETE(key). Range queries not supported. Super fast because it's just a hashmap. Used for sessions, caches, shopping carts.",
    "Challenges in Distributed RDBMS": "Horizontal scaling breaks ACID. Network partitions prevent strong consistency. Joins across nodes are expensive. Schema migrations are painful at scale. NoSQL trades some guarantees for scalability.",
    "3-Tiered Web Architecture": "Presentation layer (browser/app), Application layer (web servers/APIs), Data layer (database). Each tier can be scaled independently. Separation of concerns makes maintenance easier.",
    "XML vs JSON": "XML: verbose, tag-based, supports namespaces and schemas (XSD). JSON: lightweight, native JavaScript, easier to read, no comments. Both are tree-structured. JSON now dominates APIs.",
    "Lost Updates Problem (Concurrency)": "Thread A reads value, Thread B reads same value, Thread A writes new value, Thread B overwrites with ITS own calculation (based on old value) — Thread A's update is LOST. Fix: optimistic locking (check version) or pessimistic locking (lock the row).",
    "Historical Context: Why NoSQL?": "Late 2000s: Google Bigtable (2006), Amazon Dynamo (2007), Facebook Cassandra (2008). Web scale required trading ACID for horizontal scalability. 'NoSQL' was the answer.",
    "MongoDB ↔ RDBMS Terminology": "Database=Database. Collection=Table. Document=Row. Field=Column. Index=Index. No JOINs — embed related data or use application-level joins. Schema is flexible (different documents can have different fields).",
    "Wide Column Stores — Column Families": "Each row has a key and a set of column families. Within a family, columns are dynamic — different rows can have different columns. Stored in sorted order by row key. Great for time-series data.",
    "Key-Value Stores — RAM vs Persistent": "In-memory (Redis, Memcached): microsecond latency, data lost on restart (unless AOF/snapshot). Persistent (DynamoDB, Riak): slower but durable. Choose based on whether you can afford to lose the data.",
    "CAP Theorem — Intuitive Motivation": "If two servers can't communicate (partition), they must choose: stop responding (sacrifice availability) or respond with possibly stale data (sacrifice consistency). You CAN'T have both during a partition.",
    "Graph Databases": "Nodes = entities, edges = relationships, both can have properties. Query language: Cypher (Neo4j). Great for: social networks, fraud detection, recommendation engines where relationships are as important as entities.",

    # L10
    "Three Key Problems in Recommendation Systems": "1) Relevance (show things the user actually wants). 2) Novelty (show things they haven't seen before). 3) Serendipity (pleasant surprises). Hard to balance all three — optimising relevance alone creates filter bubbles.",
    "Utility Matrix": "A matrix where rows = users, columns = items, cells = ratings (or implicit signals like clicks). Extremely sparse — most users rate only a tiny fraction of items (Netflix: average user rates <0.1% of movies).",
    "Two Main Approaches": "Content-Based: recommend items similar to what the user liked before (based on item features). Collaborative Filtering: recommend items that similar users liked (based on user behaviour patterns). Hybrid combines both.",
    "Collaborative Filtering — Two Variants": "User-Based CF: find users similar to you, recommend what they liked. Item-Based CF: find items similar to what you liked. Item-based is more stable (item preferences change less than user preferences).",
    "Similarity Metrics": "Cosine similarity: angle between rating vectors. Pearson correlation: correlation between ratings (adjusts for rating bias). Jaccard: for implicit data (clicks, purchases). Pearson is best for explicit ratings.",
    "KNN-Based Collaborative Filtering": "Find K most similar users (or items) using cosine/Pearson similarity. Predict user's rating = weighted average of K neighbours' ratings, weighted by similarity. K is a hyperparameter.",
    "Matrix Factorization (Latent Factor Model)": "Decompose utility matrix R ≈ U × V^T. U = user factors (users × k), V = item factors (items × k). k = number of latent dimensions. Each latent factor might represent genre, mood, style — not explicitly labelled.",
    "Baseline Estimate": "Before CF, compute a baseline: b_ui = μ + b_u + b_i. μ = global mean rating, b_u = user bias (some users always rate high), b_i = item bias (some movies always rate high). Subtract before CF, add back for predictions.",
    "Evaluation Metrics": "RMSE (root mean squared error): penalises big errors more. MAE (mean absolute error): treats all errors equally. Precision@K: of top-K recommendations, how many were relevant. NDCG: ranking quality (relevant items ranked higher score more).",
    "ALS (Alternating Least Squares) — How it works": "Fix U, solve for V analytically. Fix V, solve for U analytically. Alternate until convergence. This works because fixing one set makes the problem a simple linear least-squares, which has a closed-form solution.",
    "Challenges": "Cold start: new users/items have no ratings. Data sparsity: most matrix cells are empty. Popularity bias: popular items get recommended too much. First-rater problem: a new item has no ratings yet (can't be recommended by CF).",
    "The Long Tail — Origin": "The concept that digital platforms can profitably serve many items with low individual demand, not just blockbusters. Netflix can profitably offer thousands of niche films because storage is cheap and the internet connects them to their small audiences worldwide.",
    "User-Based CF — Explicit Weighted Average Formula": "pred(u,i) = mean(r_u) + Σ[sim(u,v) × (r_v,i - mean(r_v))] / Σ|sim(u,v)|. Adjusting for each user's rating scale (their mean) before combining makes it fairer.",
    "Item-Item CF — Numerical Worked Example": "Compute similarity between item i and all other items that user u has rated. Predict = weighted average of user's ratings on similar items, weighted by similarity. More stable than user-based because item similarities change slowly.",
    "CF — Additional Limitations": "Popularity bias (over-recommend popular items), filter bubble (only shows similar content), privacy concerns (your behaviour reveals others' patterns), data sparsity issues with new users.",
    "Evaluation — Additional Metrics": "Coverage: what % of items can the system recommend. Diversity: how different are the recommendations (avoid recommending 10 nearly identical items). Novelty: recommending things the user doesn't already know.",
    "Latent Factor Model — Third Paradigm": "Beyond content-based and CF, matrix factorisation discovers LATENT (hidden) features automatically. The model figures out that 'action', 'comedy', 'mood' exist as dimensions without being told — just from rating patterns.",

    # Code section
    "T01 — Pandas Essentials": "Core Pandas: read_csv, groupby+agg, merge, pivot_table, apply. These are the foundation for any data manipulation task.",
    "T02 — Dask": "Dask mirrors Pandas API: `dd.read_csv()`, `ddf.groupby().agg()`, `ddf.compute()`. Remember: everything is lazy until `.compute()`.",
    "T03 — Spark RDD (Word Count, DNA)": "Classic RDD pattern: `sc.textFile() → flatMap(split) → map((w,1)) → reduceByKey(+) → sortByKey()`. DNA: filter by pattern, count occurrences.",
    "T04 — Spark SQL / DataFrame": "Register as temp view: `df.createOrReplaceTempView('t')`. Query: `spark.sql('SELECT ...')`. Mix DataFrame API and SQL freely.",
    "Tutorial 9 — ALS Recommender": "ALS in Spark MLlib: `ALS(userCol='userId', itemCol='movieId', ratingCol='rating').fit(train)`. `model.recommendForAllUsers(10)` gives top-10 per user.",
    "Tutorial 8 — NetworkX Graph Analytics": "NetworkX: `G = nx.Graph()`, `G.add_edges_from(...)`, `nx.betweenness_centrality(G)`, `nx.pagerank(G, alpha=0.85)`. Pure Python, great for analysis.",
    "Sparse Vector Operations (CountVectorizer output)": "CountVectorizer returns a scipy sparse matrix. Convert to dense: `.toarray()`. Access shape: `.shape`. Row = document, column = vocabulary word. Most entries are zero.",
    "UDF with Complex Return Type (ArrayType / StructType)": "Return type must be declared: `@udf(returnType=ArrayType(StringType()))`. For StructType, define schema with `StructType([StructField(...)])`. Complex types enable returning multiple values per row.",
    "Manual TF-IDF via RDD": "Compute TF as countByValue per document. Compute IDF as log(N/df) where df = number of docs containing word. Multiply: TF × IDF = final weight.",
    "Word2Vec — Nearest Words (Cosine Similarity)": "model.wv.most_similar('word', topn=10) returns 10 nearest neighbours by cosine similarity. model.wv['word'] returns the embedding vector.",
    "Tutorial 8 — NetworkX Graph Construction Exercises": "`nx.DiGraph()` for directed. `G.add_edge(u, v, weight=w)` for weighted. `nx.draw_networkx(G)` to visualise. `nx.shortest_path(G, source, target)` for path finding.",
    "Tutorial 6 — Keras Word Embedding Layer": "`Embedding(vocab_size, embed_dim, input_length=context_size)`. Input: integer-encoded word indices. Output: dense embedding vectors. Trainable — weights update during training.",
    "Tutorial 7 — PageRank (PySpark + NetworkX)": "PySpark: build edges as RDD, apply iterative rank propagation. NetworkX: `nx.pagerank(G, alpha=0.85, max_iter=100, tol=1e-6)`. Both should give same results.",
    "Spark — flatMap vs map (key difference)": "map: one input → one output. flatMap: one input → zero or more outputs (flattens the result list). `['hello world'].flatMap(str.split)` → ['hello', 'world'] not [['hello', 'world']].",
    "Exam Review — PageRank Numerical Traces": "For a 3-node graph, initialise r=[1/3,1/3,1/3], multiply by β×M, add (1-β)/3 to each entry, repeat. Track L1 norm change each iteration.",
    "Tutorial 6 — Full SparkNLP Pipeline & CBOW Training": "SparkNLP pipeline: DocumentAssembler → Tokenizer → StopWordsCleaner → Word2Vec (or CBOW). Fit pipeline on corpus, transform to get embeddings.",
    "Tutorial 7 — PageRank Formula Note": "Tutorial 7 uses: r_new = β × M × r_old + (1-β)/N × e. Where e is a vector of all ones. This is the teleportation term redistributing (1-β) of total rank equally.",
    "Tutorial 8 — Extended NetworkX Patterns": "Community detection: `nx.algorithms.community.girvan_newman(G)`. Minimum spanning tree: `nx.minimum_spanning_tree(G)`. Strongly connected: `nx.strongly_connected_components(G)`.",
    "Betweenness Centrality — Numerical Example": "For each source node s: run BFS to find all shortest paths. Count paths through each intermediate node. Divide by total paths. Sum over all source nodes. Normalise by (n-1)(n-2)/2 for undirected.",
    "Tutorial 5 - NLP: countByValue TF and distinct DF": "`rdd.countByValue()` returns dict of element→count. For TF: compute per-document then normalise. `df.distinct()` removes duplicate rows in DataFrame.",
    "Tutorial 6 - Full Vocabulary, Window Generation, Keras CBOW": "Vocabulary: unique tokens indexed 0..V-1. Window generation: for each word, take context_size words on each side. Keras CBOW: input = context indices, target = centre word index, cross-entropy loss.",
    "Tutorial 7 - PySpark PageRank 7-Step Structure": "1) Load edges. 2) Build adjacency. 3) Init ranks. 4) Compute contributions. 5) Sum contributions. 6) Apply damping. 7) Check convergence. Repeat 4-7.",
    "Tutorial 8 - Full NetworkX API Reference": "Key functions: `nx.degree_centrality`, `nx.closeness_centrality`, `nx.betweenness_centrality`, `nx.eigenvector_centrality`, `nx.clustering`, `nx.average_clustering`, `nx.diameter`.",
    "ROUGE Evaluation Metrics (Text Summarization)": "ROUGE-N: overlap of N-grams between system and reference summary. ROUGE-1: unigrams. ROUGE-2: bigrams. ROUGE-L: longest common subsequence. Higher = better summary. Like checking how much of your answer matches the model answer.",
    "T01 - Pandas Extended API": "Key Pandas: `df.describe()`, `df.info()`, `df.isnull().sum()`, `df.fillna()`, `df.drop_duplicates()`, `pd.get_dummies()`, `df.corr()`. Chain methods for concise pipelines.",
    "T04 - Extended Spark DataFrame API": "`df.withColumnRenamed('old','new')`, `df.dropDuplicates(['col'])`, `df.fillna({'col': 0})`, `df.sample(fraction=0.1)`, `df.cache()`, `df.explain()` (show query plan).",

    # Summary
    "Must-Know Formulas": "PageRank: r=βMr+(1-β)/N. TF-IDF: TF×log(N/df). Cosine similarity: A·B/(|A||B|). Jaccard: |A∩B|/|A∪B|. Levenshtein: DP with insert/delete/substitute. BPE: merge most frequent pair repeatedly.",
    "Key Comparisons at a Glance": "Map (lazy) vs Action (eager in Spark). Scale-up vs Scale-out. ACID vs BASE. Strong vs Eventual consistency. Stemming vs Lemmatisation. Skip-gram vs CBOW. Content-based vs CF.",
    "CAP Theorem Cheat Sheet": "CP systems: MongoDB, HBase, Zookeeper (consistent during partition, may be unavailable). AP systems: Cassandra, CouchDB, DynamoDB (available during partition, may be inconsistent). CA: traditional RDBMS (no partition tolerance — only works on single node).",
    "PageRank Algorithm Steps": "1) Init r=[1/N]. 2) r_new = β×M×r + (1-β)/N. 3) Handle dead-ends (add leaked rank). 4) Compute L1 change. 5) If change < threshold: stop. Else: r=r_new, goto 2.",
    "Text Preprocessing Checklist": "1) Lowercase. 2) Remove punctuation/HTML. 3) Tokenise. 4) Remove stopwords. 5) Stem OR lemmatise. 6) Vectorise (BoW/TF-IDF/embeddings). 7) Handle OOV (BPE or [UNK] token).",
}

# ── helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_toc_entry(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 * (level - 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    if level == 1:
        run.bold = True

def add_eli5(doc, text):
    if not text:
        return
    p = doc.add_paragraph()
    run_label = p.add_run('ELI5: ')
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    run_label.font.size = Pt(10)
    run_content = p.add_run(text)
    run_content.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    run_content.font.size = Pt(10)
    run_content.italic = True
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)

def parse_inline(para, el):
    """Add inline text from an element's children to a docx paragraph."""
    for node in el.children:
        if isinstance(node, NavigableString):
            txt = str(node)
            if txt.strip() or ' ' in txt:
                run = para.add_run(txt)
                run.font.size = Pt(10.5)
        elif isinstance(node, Tag):
            n = node.name
            txt = node.get_text()
            run = para.add_run(txt)
            run.font.size = Pt(10.5)
            if n in ('strong', 'b'):
                run.bold = True
            elif n in ('em', 'i'):
                run.italic = True
            elif n == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

def add_element(doc, el):
    """Recursively add an HTML element to the docx."""
    if isinstance(el, NavigableString):
        text = str(el).strip()
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(2)
        return

    tag = el.name
    classes = el.get('class', [])

    if tag == 'h3':
        h3_text = el.get_text(strip=True)
        doc.add_heading(h3_text, level=3)
        eli5_text = ELI5.get(h3_text, '')
        add_eli5(doc, eli5_text)

    elif tag == 'h4':
        doc.add_heading(el.get_text(strip=True), level=4)

    elif tag == 'table':
        rows = el.find_all('tr')
        if not rows:
            return
        # find max cols
        max_cols = max(len(r.find_all(['td','th'])) for r in rows)
        if max_cols == 0:
            return
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        for ri, row in enumerate(rows):
            cells = row.find_all(['td','th'])
            for ci, cell in enumerate(cells):
                if ci >= max_cols:
                    break
                tc = table.rows[ri].cells[ci]
                tc.text = ''
                p = tc.paragraphs[0]
                is_header = (cell.name == 'th')
                parse_inline(p, cell)
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    if is_header:
                        run.bold = True
                if is_header:
                    set_cell_bg(tc, 'D1D5DB')
        doc.add_paragraph()

    elif tag in ('ul', 'ol'):
        items = el.find_all('li', recursive=False)
        if not items:
            # try all li
            items = el.find_all('li')
        for item in items:
            style = 'List Bullet' if tag == 'ul' else 'List Number'
            p = doc.add_paragraph(style=style)
            parse_inline(p, item)
            for run in p.runs:
                run.font.size = Pt(10.5)

    elif tag == 'pre' or (tag == 'div' and ('formula' in classes or 'code-block' in classes)):
        code_el = el.find('code') or el
        code_text = code_el.get_text()
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4) if 'formula' in classes else RGBColor(0xA7, 0x8B, 0xFA)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

    elif tag == 'div' and 'callout' in classes:
        text = el.get_text(strip=True)
        p = doc.add_paragraph()
        run_icon = p.add_run('NOTE: ')
        run_icon.bold = True
        run_icon.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
        run_content = p.add_run(text)
        run_content.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
        run_content.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)

    elif tag == 'div' and 'step' in classes:
        step_num_el = el.find(class_='step-num')
        step_num = step_num_el.get_text(strip=True) if step_num_el else ''
        if step_num_el:
            step_num_el.decompose()
        text = el.get_text(strip=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        run_n = p.add_run(f'Step {step_num}: ' if step_num else '')
        run_n.bold = True
        run_n.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        run_t = p.add_run(text)
        run_t.font.size = Pt(10.5)

    elif tag == 'p':
        p = doc.add_paragraph()
        parse_inline(p, el)
        p.paragraph_format.space_after = Pt(4)

    elif tag in ('div', 'section', 'article', 'span'):
        # recurse into container elements
        for child in el.children:
            add_element(doc, child)

    elif tag in ('br',):
        pass  # skip
    elif tag in ('code',):
        p = doc.add_paragraph()
        run = p.add_run(el.get_text())
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
    else:
        # fallback: just add text
        text = el.get_text(strip=True)
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(2)

# ── Section metadata ──────────────────────────────────────────────────────────
SECTIONS = [
    ('l1',  'L1: Introduction to Big Data Analytics'),
    ('l2',  'L2: The Path to Parallelism (Dask, MapReduce, HDFS)'),
    ('l3',  'L3: Apache Spark'),
    ('l4',  'L4: Hive, Shark & SparkSQL'),
    ('l5',  'L5: Intro to Natural Language Processing (NLP)'),
    ('l6',  'L6: Large Language Models & Word Embeddings'),
    ('l7',  'L7: Link Analysis — PageRank'),
    ('l8',  'L8: Graph Analytics'),
    ('l9',  'L9: Big Data Storage & NoSQL'),
    ('l10', 'L10: Recommendation Systems'),
    ('code','Tutorial Code Reference (T01–T09 & Notebooks)'),
    ('summary', 'Quick Reference Summary'),
]

# ── Build document ────────────────────────────────────────────────────────────
html_text = HTML.read_text(encoding='utf-8')
soup = BeautifulSoup(html_text, 'html.parser')

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# Title
title_p = doc.add_heading('DSAI4205 Big Data Analytics — Full Exam Notes (with ELI5)', level=0)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('All content from L1–L10 + Tutorial Code Reference + Quick Reference Summary')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
sub.runs[0].font.size = Pt(11)

eli5_note = doc.add_paragraph()
r1 = eli5_note.add_run('Green italic text = ELI5 (Explain Like I\'m 5) — plain-English analogy for every concept')
r1.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
r1.italic = True
r1.font.size = Pt(10)
eli5_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# TOC
doc.add_heading('Table of Contents', level=1)
for sid, title in SECTIONS:
    add_toc_entry(doc, title, level=1)
    body_div = soup.find('div', id=f'body-{sid}')
    if body_div:
        for h3 in body_div.find_all('h3', recursive=False):
            add_toc_entry(doc, h3.get_text(strip=True), level=2)
        # also get h3 inside direct children divs
        for child_div in body_div.find_all('div', recursive=False):
            for h3 in child_div.find_all('h3', recursive=False):
                add_toc_entry(doc, h3.get_text(strip=True), level=2)

doc.add_page_break()

# Main content
for sid, section_title in SECTIONS:
    print(f'Processing {sid}...')
    doc.add_heading(section_title, level=1)

    body_div = soup.find('div', id=f'body-{sid}')
    if not body_div:
        p = doc.add_paragraph(f'[Section {sid} not found in HTML]')
        p.runs[0].font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
        doc.add_page_break()
        continue

    for child in body_div.children:
        add_element(doc, child)

    doc.add_page_break()

doc.save(str(DOCX))
print(f'\nSaved: {DOCX}')

# Verify
from docx import Document as DDoc
d2 = DDoc(str(DOCX))
non_empty = [p for p in d2.paragraphs if p.text.strip()]
tables = d2.tables
h2s = [p for p in d2.paragraphs if p.style.name.startswith('Heading 1')]
h3s = [p for p in d2.paragraphs if p.style.name.startswith('Heading 3')]
eli5s = [p for p in d2.paragraphs if p.text.startswith('ELI5:')]
print(f'Non-empty paragraphs: {len(non_empty)}')
print(f'Tables: {len(tables)}')
print(f'H1 sections: {len(h2s)}')
print(f'H3 subsections: {len(h3s)}')
print(f'ELI5 entries: {len(eli5s)}')
